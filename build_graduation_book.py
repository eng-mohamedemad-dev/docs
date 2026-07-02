"""
Build the editable CrimeLens graduation book (front matter + Chapter One).

Output:
    CrimeLens_Graduation_Book.docx

The document is intentionally generated from structured content so later
chapters can reuse the exact same typography, page system, captions, tables,
headers, footers, and IEEE-style references.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor

from chapter_three_content import (
    CHAPTER_THREE_FIGURES,
    CHAPTER_THREE_REFERENCES,
    CHAPTER_THREE_TABLES,
    chapter_three_page_specs,
)
from chapter_four_content import (
    CHAPTER_FOUR_FIGURES,
    CHAPTER_FOUR_REFERENCES,
    CHAPTER_FOUR_TABLES,
    chapter_four_page_specs,
)
from chapter_five_content import (
    CHAPTER_FIVE_FIGURES,
    CHAPTER_FIVE_REFERENCES,
    CHAPTER_FIVE_TABLES,
    chapter_five_page_specs,
)
from chapter_six_content import (
    CHAPTER_SIX_FIGURES,
    CHAPTER_SIX_REFERENCES,
    CHAPTER_SIX_TABLES,
    chapter_six_page_specs,
)
from chapter_seven_content import (
    CHAPTER_SEVEN_FIGURES,
    CHAPTER_SEVEN_REFERENCES,
    CHAPTER_SEVEN_TABLES,
    chapter_seven_page_specs,
)


ROOT = Path(__file__).resolve().parents[2]
HERE = Path(__file__).resolve().parent
OUTPUT = HERE / "CrimeLens_Graduation_Book.docx"

PRESENTATION_ASSETS = HERE / "assets"
BRANDING = PRESENTATION_ASSETS / "branding"
TEAM_ASSETS = PRESENTATION_ASSETS / "team"
POSTER_IMAGES = PRESENTATION_ASSETS / "poster" / "images"
BROCHURE_IMAGES = PRESENTATION_ASSETS / "brochure" / "images"
LEGACY_PRESENTATION_ASSETS = ROOT / "project" / "presentation" / "assets"
LEGACY_BRANDING = LEGACY_PRESENTATION_ASSETS / "branding"
LEGACY_TEAM_ASSETS = LEGACY_PRESENTATION_ASSETS / "team"
LEGACY_POSTER_IMAGES = HERE.parent / "poster" / "images"
LEGACY_BROCHURE_IMAGES = HERE.parent / "brochure" / "images"
COVERS = HERE / "covers"
CHAPTER_IMAGES = HERE / "chapter-images"

NAVY = "07111F"
NAVY_2 = "0E1B31"
NAVY_3 = "172640"
CYAN = "38BDF8"
CYAN_DARK = "0E7490"
GOLD = "D4AF37"
GREEN = "22C55E"
PURPLE = "A855F7"
RED = "EF4444"
ORANGE = "F59E0B"
BLUE = "2563EB"
TEAL = "2DD4BF"
TEXT = "162033"
MUTED = "5B6B80"
LIGHT = "F4F7FB"
WHITE = "FFFFFF"
BLACK = "000000"
INK_MUTED = "555555"
BORDER = "C9D5E5"

CONTENT_FONT_BUMP = 1
CONTENT_FONT_BUMP_MAX_SIZE = 13.0

EDITORIAL_RULES = (
    "Use a white academic page background throughout the entire book.",
    "Use color only for restrained headings, rules, borders, and small accents.",
    "Define every abbreviation at its first use as Full Term (ABBR).",
    "After its first definition, use the abbreviation alone within the same chapter.",
    "Do not create a standalone list of abbreviations.",
)

PROJECT_TITLE = "CrimeLens: AI-Assisted Computer-Aided Dispatch and Smart Surveillance System"
UNIVERSITY = "Beni-Suef University"
FACULTY = "Faculty of Computers and Artificial Intelligence"
DEPARTMENT = "Computer Science Department"
ACADEMIC_YEAR = "2025 / 2026"

SUPERVISORS = [
    ("Dr. Mohamed Salah Reda", "Project Supervisor"),
    ("Dr. Ahmed El-Naggar", "Project Supervisor"),
    ("Osama Hefny", "Teaching Assistant"),
]

TEAM = [
    ("Mohamed Emad El Deen", "220357", "Full Stack / Team Lead", "team-220357.png"),
    ("Mohamed Mahmoud Alam", "220366", "Mobile — Flutter", "team-220366.png"),
    ("Abdalla Mohamed Swesy", "220239", "Mobile — Flutter", "team-220239.png"),
    ("Ahmed Mohamed Sadawu", "220051", "AI / ML Engineer", "team-220051.png"),
    ("Abdelrahman Amer Sayed", "220223", "DevOps & Streaming", "team-220223.png"),
    ("Ahmed Salama Ahmed", "220031", "QA & Testing", "team-220031.png"),
    ("Mohamed Ahmed Atia", "210270", "UI / UX Designer", "team-210270.png"),
]


BRANDING_FALLBACKS = {
    "university-logo.jpeg": ("logo-university.jpeg", "logo-university-clean.png"),
    "faculty-logo.jpg": ("logo-faculty.jpg", "logo-faculty-clean.png"),
    "crimelens-logo.png": ("logo-crimelens.png", "logo-keyed.png", "logo-crimelens2.png"),
}


def first_existing(*paths: Path) -> Path:
    for path in paths:
        if path.exists():
            return path
    return paths[0]


def branding_asset(filename: str) -> Path:
    candidates = [BRANDING / filename, LEGACY_BRANDING / filename]
    for fallback_name in BRANDING_FALLBACKS.get(filename, (filename,)):
        candidates.extend([
            POSTER_IMAGES / fallback_name,
            BROCHURE_IMAGES / fallback_name,
            LEGACY_POSTER_IMAGES / fallback_name,
            LEGACY_BROCHURE_IMAGES / fallback_name,
        ])
    return first_existing(*candidates)


def team_asset(filename: str) -> Path:
    candidates = [
        TEAM_ASSETS / filename,
        POSTER_IMAGES / "team" / filename,
        BROCHURE_IMAGES / "team" / filename,
        LEGACY_TEAM_ASSETS / filename,
        LEGACY_POSTER_IMAGES / "team" / filename,
        LEGACY_BROCHURE_IMAGES / "team" / filename,
    ]
    stem = Path(filename).stem
    for extension in (".png", ".jpg", ".jpeg"):
        alternate = f"{stem}{extension}"
        candidates.extend([
            TEAM_ASSETS / alternate,
            POSTER_IMAGES / "team" / alternate,
            BROCHURE_IMAGES / "team" / alternate,
            LEGACY_TEAM_ASSETS / alternate,
            LEGACY_POSTER_IMAGES / "team" / alternate,
            LEGACY_BROCHURE_IMAGES / "team" / alternate,
        ])
    return first_existing(*candidates)


FIGURES = [
    ("Figure 1.1", "Chapter-opening command-center background", 1),
    ("Figure 1.2", "Evolution from passive CCTV monitoring to an AI-assisted CAD workflow", 3),
    ("Figure 1.3", "CrimeLens high-level operational workflow", 7),
    ("Figure 1.4", "CrimeLens objective hierarchy", 11),
    ("Figure 1.5", "Project scope boundary", 12),
    ("Figure 1.6", "Iterative development methodology", 14),
    ("Figure 1.7", "Five-month work plan", 15),
    ("Figure 1.8", "Market-category positioning map", 17),
]

TABLES = [
    ("Table 1.1", "Major video anomaly-detection model families", 5),
    ("Table 1.2", "Benchmark datasets relevant to surveillance anomaly detection", 6),
    ("Table 1.3", "Comparison of supported streaming protocols", 7),
    ("Table 1.4", "Problem statement dimensions", 10),
    ("Table 1.5", "Stakeholders and operational pain points", 11),
    ("Table 1.6", "General and specific project objectives", 14),
    ("Table 1.7", "In-scope and out-of-scope boundaries", 16),
    ("Table 1.8", "Methodology phases and outputs", 19),
    ("Table 1.9", "High-level comparison of solution categories", 23),
    ("Table 1.10", "Book chapters and their purpose", 25),
]

CHAPTER_TWO_FIGURES = [
    ("Figure 2.1", "Chapter-opening literature-review background", 28),
    ("Figure 2.2", "Evolution of surveillance-video analysis methods", 30),
    ("Figure 2.3", "Contextual taxonomy of video anomalies", 31),
    ("Figure 2.4", "Supervision paradigms for video anomaly detection", 32),
    ("Figure 2.5", "Spatial, temporal, and object feature representations", 33),
    ("Figure 2.6", "Literature-selection and evidence-synthesis workflow", 34),
    ("Figure 2.7", "Taxonomy of algorithm families reviewed in Chapter Two", 35),
    ("Figure 2.8", "Two-dimensional convolutional classification pipeline", 36),
    ("Figure 2.9", "Two-stream appearance and optical-flow architecture", 37),
    ("Figure 2.10", "Three-dimensional convolution over space and time", 38),
    ("Figure 2.11", "Recurrent temporal-modeling pipeline", 39),
    ("Figure 2.12", "Reconstruction-based autoencoder anomaly scoring", 40),
    ("Figure 2.13", "Future-frame prediction and prediction-error scoring", 41),
    ("Figure 2.14", "Adversarial learning for video anomaly detection", 42),
    ("Figure 2.15", "Multiple-instance learning with video bags and temporal instances", 43),
    ("Figure 2.16", "Robust temporal feature-magnitude learning", 44),
    ("Figure 2.17", "Object-detection and scene-evidence pipeline", 45),
    ("Figure 2.18", "Space-time attention for video understanding", 46),
    ("Figure 2.19", "Hybrid evidence-fusion strategies", 48),
    ("Figure 2.20", "Evaluation metrics and threshold behavior", 50),
    ("Figure 2.21", "Dataset landscape by realism, duration, and annotation level", 51),
    ("Figure 2.22", "UCF-Crime categories and weak labels", 52),
    ("Figure 2.23", "Representative compact anomaly-detection benchmarks", 53),
    ("Figure 2.24", "RWF-2000 violence dataset characteristics", 54),
    ("Figure 2.25", "Sources of cross-dataset domain shift", 55),
    ("Figure 2.26", "Chronological evolution of representative literature", 56),
    ("Figure 2.27", "Research-gap and opportunity map", 57),
]

CHAPTER_TWO_TABLES = [
    ("Table 2.1", "Scope and boundaries of the literature review", 29),
    ("Table 2.2", "Categories of video anomalies", 31),
    ("Table 2.3", "Comparison of supervision paradigms", 32),
    ("Table 2.4", "Feature representations used in surveillance analysis", 33),
    ("Table 2.5", "Algorithm families and principal trade-offs", 35),
    ("Table 2.6", "Comparison of two-dimensional and three-dimensional convolution", 38),
    ("Table 2.7", "Reconstruction, prediction, and adversarial approaches", 42),
    ("Table 2.8", "Representative benchmark datasets", 51),
    ("Table 2.9", "Evaluation metrics and interpretation", 50),
    ("Table 2.10", "Reported results selected from primary sources", 56),
    ("Table 2.11", "Literature-survey synthesis", 57),
]

ALL_FIGURES = (
    FIGURES
    + CHAPTER_TWO_FIGURES
    + CHAPTER_THREE_FIGURES
    + CHAPTER_FOUR_FIGURES
    + CHAPTER_FIVE_FIGURES
    + CHAPTER_SIX_FIGURES
    + CHAPTER_SEVEN_FIGURES
)
ALL_TABLES = (
    TABLES
    + CHAPTER_TWO_TABLES
    + CHAPTER_THREE_TABLES
    + CHAPTER_FOUR_TABLES
    + CHAPTER_FIVE_TABLES
    + CHAPTER_SIX_TABLES
    + CHAPTER_SEVEN_TABLES
)

ABBREVIATION_DEFINITIONS = [
    ("2D", "Two-Dimensional"),
    ("3D", "Three-Dimensional"),
    ("AE", "Autoencoder"),
    ("AES", "Advanced Encryption Standard"),
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("AUC", "Area Under the Receiver Operating Characteristic Curve"),
    ("BOLO", "Be On the Lookout"),
    ("CAD", "Computer-Aided Dispatch"),
    ("CCTV", "Closed-Circuit Television"),
    ("CI/CD", "Continuous Integration and Continuous Deployment"),
    ("CNN", "Convolutional Neural Network"),
    ("CPU", "Central Processing Unit"),
    ("CRUD", "Create, Read, Update, and Delete"),
    ("CSRF", "Cross-Site Request Forgery"),
    ("CV", "Computer Vision"),
    ("DFD", "Data Flow Diagram"),
    ("DTO", "Data Transfer Object"),
    ("E2E", "End-to-End"),
    ("EER", "Equal Error Rate"),
    ("ERD", "Entity–Relationship Diagram"),
    ("F-AUC", "Frame-Level Area Under the Curve"),
    ("FCM", "Firebase Cloud Messaging"),
    ("FPS", "Frames Per Second"),
    ("GAN", "Generative Adversarial Network"),
    ("GEO", "Geospatial"),
    ("GIS", "Geographic Information System"),
    ("GPS", "Global Positioning System"),
    ("GPU", "Graphics Processing Unit"),
    ("HLS", "HTTP Live Streaming"),
    ("HMAC", "Hash-Based Message Authentication Code"),
    ("HTTP", "Hypertext Transfer Protocol"),
    ("HTTPS", "Hypertext Transfer Protocol Secure"),
    ("IP", "Internet Protocol"),
    ("JSON", "JavaScript Object Notation"),
    ("LAN", "Local Area Network"),
    ("LSTM", "Long Short-Term Memory"),
    ("MAC", "Message Authentication Code"),
    ("MAE", "Mean Absolute Error"),
    ("MIL", "Multiple Instance Learning"),
    ("MIME", "Multipurpose Internet Mail Extensions"),
    ("ML", "Machine Learning"),
    ("MLOps", "Machine Learning Operations"),
    ("MSE", "Mean Squared Error"),
    ("NAT", "Network Address Translation"),
    ("NFR", "Non-Functional Requirement"),
    ("ONVIF", "Open Network Video Interface Forum"),
    ("P-AUC", "Pixel-Level Area Under the Curve"),
    ("PTZ", "Pan, Tilt, and Zoom"),
    ("QA", "Quality Assurance"),
    ("RBAC", "Role-Based Access Control"),
    ("REST", "Representational State Transfer"),
    ("RFC", "Request for Comments"),
    ("ROC", "Receiver Operating Characteristic"),
    ("RTSP", "Real-Time Streaming Protocol"),
    ("SDK", "Software Development Kit"),
    ("SMS", "Short Message Service"),
    ("SOS", "Emergency Distress Signal"),
    ("SQL", "Structured Query Language"),
    ("SRS", "Software Requirements Specification"),
    ("TLS", "Transport Layer Security"),
    ("TTL", "Time to Live"),
    ("UCF", "University of Central Florida"),
    ("UI", "User Interface"),
    ("UML", "Unified Modeling Language"),
    ("URL", "Uniform Resource Locator"),
    ("UX", "User Experience"),
    ("VMS", "Video Management System"),
    ("VPN", "Virtual Private Network"),
    ("WebRTC", "Web Real-Time Communication"),
    ("YOLO", "You Only Look Once"),
]

DEFINED_ABBREVIATIONS: set[str] = set()


REFERENCES = [
    (
        "[1]",
        'M. H. Sharif, L. Jiao, and C. W. Omlin, "Deep crowd anomaly detection: '
        'state-of-the-art, challenges, and future research directions," Artificial '
        "Intelligence Review, vol. 58, art. no. 139, 2025, "
        "doi: 10.1007/s10462-024-11092-8.",
    ),
    (
        "[2]",
        'W. Sultani, C. Chen, and M. Shah, "Real-world anomaly detection in '
        'surveillance videos," in Proc. IEEE/CVF Conference on Computer Vision '
        "and Pattern Recognition (CVPR), 2018, pp. 6479–6488.",
    ),
    (
        "[3]",
        'H. Schulzrinne, A. Rao, and R. Lanphier, "Real-Time Streaming Protocol '
        'Version 2.0," Internet Engineering Task Force, RFC 7826, Dec. 2016.',
    ),
    (
        "[4]",
        'R. Pantos and W. May, "HTTP Live Streaming," Internet Engineering Task '
        "Force, RFC 8216, Aug. 2017.",
    ),
    (
        "[5]",
        'World Wide Web Consortium, "WebRTC 1.0: Real-Time Communication Between '
        'Browsers," W3C Recommendation, Jan. 2021.',
    ),
    (
        "[6]",
        'Genetec Inc., "Security Center unified security platform," official '
        "product information, accessed Jun. 24, 2026. [Online]. Available: "
        "https://www.genetec.com/products/unified-security/security-center",
    ),
    (
        "[7]",
        'Motorola Solutions, "Command center software," official product '
        "information portal, accessed Jun. 24, 2026. [Online]. Available: "
        "https://www.motorolasolutions.com/",
    ),
    (
        "[8]",
        'Axon Enterprise, Inc., "Public safety and real-time operations product '
        'information," accessed Jun. 24, 2026. [Online]. Available: '
        "https://www.axon.com/",
    ),
    (
        "[9]",
        'Verkada Inc., "Cloud-managed video security product information," '
        "accessed Jun. 24, 2026. [Online]. Available: https://www.verkada.com/",
    ),
    (
        "[10]",
        "CrimeLens Project Team, CrimeLens internal architecture, workflow, "
        "requirements, security, AI-integration, and testing documentation, "
        "Beni-Suef University, 2026.",
    ),
]

CHAPTER_TWO_REFERENCES = [
    (
        "[11]",
        'M. Hasan, J. Choi, J. Neumann, A. K. Roy-Chowdhury, and L. S. Davis, '
        '"Learning temporal regularity in video sequences," in Proc. IEEE Conference '
        "on Computer Vision and Pattern Recognition (CVPR), 2016, pp. 733–742.",
    ),
    (
        "[12]",
        'W. Liu, W. Luo, D. Lian, and S. Gao, "Future frame prediction for anomaly '
        'detection—A new baseline," in Proc. IEEE Conference on Computer Vision and '
        "Pattern Recognition (CVPR), 2018, pp. 6536–6545.",
    ),
    (
        "[13]",
        'D. Tran, L. Bourdev, R. Fergus, L. Torresani, and M. Paluri, "Learning '
        'spatiotemporal features with 3D convolutional networks," in Proc. IEEE '
        "International Conference on Computer Vision (ICCV), 2015, pp. 4489–4497.",
    ),
    (
        "[14]",
        'J. Carreira and A. Zisserman, "Quo Vadis, action recognition? A new model '
        'and the Kinetics dataset," in Proc. IEEE Conference on Computer Vision and '
        "Pattern Recognition (CVPR), 2017, pp. 6299–6308.",
    ),
    (
        "[15]",
        'K. Simonyan and A. Zisserman, "Two-stream convolutional networks for action '
        'recognition in videos," in Advances in Neural Information Processing Systems, '
        "vol. 27, 2014.",
    ),
    (
        "[16]",
        'J. Redmon, S. Divvala, R. Girshick, and A. Farhadi, "You only look once: '
        'Unified, real-time object detection," in Proc. IEEE Conference on Computer '
        "Vision and Pattern Recognition (CVPR), 2016, pp. 779–788.",
    ),
    (
        "[17]",
        'I. Goodfellow et al., "Generative adversarial nets," in Advances in Neural '
        "Information Processing Systems, vol. 27, 2014.",
    ),
    (
        "[18]",
        'M. Cheng, K. Cai, and M. Li, "RWF-2000: An open large scale video database '
        'for violence detection," in Proc. International Conference on Pattern '
        "Recognition (ICPR), 2020.",
    ),
    (
        "[19]",
        'Y. Tian, G. Pang, Y. Chen, R. Singh, J. W. Verjans, and G. Carneiro, '
        '"Weakly-supervised video anomaly detection with robust temporal feature '
        'magnitude learning," in Proc. IEEE International Conference on Computer '
        "Vision (ICCV), 2021.",
    ),
    (
        "[20]",
        'G. Bertasius, H. Wang, and L. Torresani, "Is space-time attention all you '
        'need for video understanding?" in Proc. International Conference on Machine '
        "Learning (ICML), 2021.",
    ),
    (
        "[21]",
        'A. Arnab, M. Dehghani, G. Heigold, C. Sun, M. Lučić, and C. Schmid, '
        '"ViViT: A video vision transformer," in Proc. IEEE International Conference '
        "on Computer Vision (ICCV), 2021, pp. 6836–6846.",
    ),
    (
        "[22]",
        'Y. Lu, M. K. K. Reddy, S. S. Nabavi, and Y. Wang, "Future frame prediction '
        'using convolutional VRNN for anomaly detection," in Proc. IEEE Winter '
        "Conference on Applications of Computer Vision (WACV), 2020.",
    ),
]


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def reset_abbreviation_tracking() -> None:
    DEFINED_ABBREVIATIONS.clear()


def expand_abbreviations(text: str) -> str:
    """Define each abbreviation once in prose, then retain the short form."""
    expanded_text = text

    for abbreviation, full_term in ABBREVIATION_DEFINITIONS:
        existing_definition = f"{full_term} ({abbreviation})"
        if existing_definition in expanded_text:
            DEFINED_ABBREVIATIONS.add(abbreviation)
            for nested_abbreviation, _ in ABBREVIATION_DEFINITIONS:
                if re.search(
                    rf"(?<!\w){re.escape(nested_abbreviation)}(?!\w)",
                    full_term,
                ):
                    DEFINED_ABBREVIATIONS.add(nested_abbreviation)
            continue

        if abbreviation in DEFINED_ABBREVIATIONS:
            continue

        pattern = rf"(?<!\w){re.escape(abbreviation)}(?!\w)"
        if re.search(pattern, expanded_text):
            expanded_text = re.sub(
                pattern,
                existing_definition,
                expanded_text,
                count=1,
            )
            DEFINED_ABBREVIATIONS.add(abbreviation)

    return expanded_text


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_table_borders(table, color: str = BORDER, size: int = 6) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": size, "color": color},
                bottom={"val": "single", "sz": size, "color": color},
                left={"val": "single", "sz": size, "color": color},
                right={"val": "single", "sz": size, "color": color},
                insideH={"val": "single", "sz": size, "color": color},
                insideV={"val": "single", "sz": size, "color": color},
            )


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top: int = 100, start: int = 120,
                     bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top, "start": start, "bottom": bottom, "end": end
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_page_number_format(section, fmt: str, start: int | None = None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)
    start_attr = qn("w:start")
    if start is None:
        if start_attr in pg_num_type.attrib:
            del pg_num_type.attrib[start_attr]
    else:
        pg_num_type.set(start_attr, str(start))


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def enable_update_fields(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def set_document_background(document: Document, color: str = WHITE) -> None:
    background = OxmlElement("w:background")
    background.set(qn("w:color"), color)
    document._element.insert(0, background)


def set_run_font(run, name: str, size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def setup_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Cambria"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
    normal.font.size = Pt(11.5 + CONTENT_FONT_BUMP)
    normal.font.color.rgb = rgb(TEXT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.space_after = Pt(5)

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    title.font.size = Pt(31)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_after = Pt(12)

    for name, size, color, before, after in [
        ("Heading 1", 22, NAVY, 10, 8),
        ("Heading 2", 16, CYAN_DARK, 9, 5),
        ("Heading 3", 12.5, NAVY_3, 7, 4),
        ("Heading 4", 11, PURPLE, 6, 3),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name, color in [
        ("Figure Caption", CYAN_DARK),
        ("Table Caption", GOLD),
    ]:
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[style_name]
        style.font.name = "Cambria"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
        style.font.size = Pt(9.5 + CONTENT_FONT_BUMP)
        style.font.italic = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.keep_with_next = True

    if "Book Subtitle" not in styles:
        subtitle = styles.add_style("Book Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = styles["Book Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    subtitle.font.size = Pt(12 + CONTENT_FONT_BUMP)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)

    if "Front Matter Title" not in styles:
        fm_title = styles.add_style("Front Matter Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        fm_title = styles["Front Matter Title"]
    fm_title.font.name = "Aptos Display"
    fm_title._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    fm_title.font.size = Pt(24)
    fm_title.font.bold = True
    fm_title.font.color.rgb = rgb(NAVY)
    fm_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fm_title.paragraph_format.space_after = Pt(18)

    if "Chapter Kicker" not in styles:
        kicker = styles.add_style("Chapter Kicker", WD_STYLE_TYPE.PARAGRAPH)
    else:
        kicker = styles["Chapter Kicker"]
    kicker.font.name = "Aptos"
    kicker._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    kicker.font.size = Pt(9 + CONTENT_FONT_BUMP)
    kicker.font.bold = True
    kicker.font.color.rgb = rgb(CYAN_DARK)
    kicker.paragraph_format.space_after = Pt(2)
    kicker.paragraph_format.keep_with_next = True


def add_page_border(section, color: str = BLACK, sz: int = 14, space: int = 24) -> None:
    """Add a double border frame around the whole page (every page of the section)."""
    sect_pr = section._sectPr
    for existing in sect_pr.findall(qn("w:pgBorders")):
        sect_pr.remove(existing)
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "double")
        edge.set(qn("w:sz"), str(sz))
        edge.set(qn("w:space"), str(space))
        edge.set(qn("w:color"), color)
        pg_borders.append(edge)
    pg_mar = sect_pr.find(qn("w:pgMar"))
    if pg_mar is not None:
        pg_mar.addnext(pg_borders)
    else:
        sect_pr.append(pg_borders)


def configure_a4(section, top: float = 1.9, bottom: float = 1.8,
                 left: float = 2.15, right: float = 1.9) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    add_page_border(section)


def unlink_section(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False


def add_header(section, label: str, text_color: str = MUTED, accent: str = CYAN) -> None:
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(label.upper())
    set_run_font(run, "Aptos", 8.5, text_color, True)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), accent)
    borders.append(bottom)
    p_pr.append(borders)


def add_footer(section, roman: bool = False, mono: bool = False) -> None:
    crime_color = BLACK if mono else CYAN
    lens_color = BLACK if mono else GOLD
    muted_color = BLACK if mono else MUTED
    rule_color = BLACK if mono else BORDER
    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(1, 3, Inches(6.9))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(2.4), Inches(2.1), Inches(2.4)]
    for index, width in enumerate(widths):
        table.columns[index].width = width
        table.cell(0, index).width = width
    left = table.cell(0, 0).paragraphs[0]
    left_run = left.add_run("Crime")
    set_run_font(left_run, "Aptos", 8.5, crime_color, True)
    right_run = left.add_run("Lens")
    set_run_font(right_run, "Aptos", 8.5, lens_color, True)
    center = table.cell(0, 1).paragraphs[0]
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_run = center.add_run("GRADUATION PROJECT")
    set_run_font(center_run, "Aptos", 8, muted_color, True)
    right = table.cell(0, 2).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(right, "PAGE", "1")
    set_run_font(right.runs[-1], "Aptos", 15, muted_color, True)
    for cell in table.rows[0].cells:
        set_cell_border(
            cell,
            top={"val": "single", "sz": 8, "color": rule_color},
            bottom={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
        )


def add_scope_footer(section, scope_label: str) -> None:
    """Print-ready footer with the scope title and plain black page number.
    """
    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].paragraph_format.space_after = Pt(0)

    table = footer.add_table(1, 2, Inches(6.9))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(5.85), Inches(1.05)]
    for index, width in enumerate(widths):
        table.columns[index].width = width
        table.cell(0, index).width = width

    left = table.cell(0, 0)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    set_cell_margins(left, 0, 60, 0, 90)
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run(scope_label)
    set_run_font(lr, "Aptos", 11.5, BLACK, True)
    set_cell_border(
        left,
        bottom={"val": "single", "sz": 22, "color": BLACK},
        top={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"},
    )

    right = table.cell(0, 1)
    set_cell_shading(right, WHITE)
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(right, 40, 60, 40, 60)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_run = rp.add_run()
    set_run_font(page_run, "Aptos", 15, BLACK, True)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    page_run._r.extend([begin, instr, separate, text, end])


_MONO_BORDER_TAGS = {
    qn(f"w:{edge}")
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV", "start", "end", "bar", "between")
}


def monochrome_book_body(document: Document) -> None:
    """Convert the whole document body to pure black & white for print, skipping the
    first table (the navy + gold book cover, which is intentionally left in colour).
    """
    cover_skipped = False
    for child in list(document.element.body):
        if not cover_skipped and child.tag == qn("w:tbl"):
            cover_skipped = True
            continue
        monochrome_elements([child])


def increase_body_content_font_sizes(document: Document) -> None:
    """Increase readable page content only; headers and footers are separate parts."""
    for tag in ("w:sz", "w:szCs"):
        for node in document.element.body.iter(qn(tag)):
            raw_value = node.get(qn("w:val"))
            if not raw_value:
                continue
            try:
                half_points = int(raw_value)
            except ValueError:
                continue
            point_size = half_points / 2
            if point_size <= CONTENT_FONT_BUMP_MAX_SIZE:
                node.set(
                    qn("w:val"),
                    str(int(round((point_size + CONTENT_FONT_BUMP) * 2))),
                )


def monochrome_elements(elements) -> None:
    """Force a set of body elements to pure black & white for print: every text and
    border colour becomes black, and every coloured cell/paragraph fill becomes white.
    Raster images (e.g. chapter banner photos) are left untouched.
    """
    for el in elements:
        # Force every run to black so style-defined colours are overridden too.
        for run in el.iter(qn("w:r")):
            rpr = run.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                run.insert(0, rpr)
            color = rpr.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                rpr.append(color)
            color.set(qn("w:val"), "000000")
        # Any remaining colour references (e.g. paragraph marks) → black.
        for color in el.iter(qn("w:color")):
            color.set(qn("w:val"), "000000")
        # Coloured fills → white.
        for shd in el.iter(qn("w:shd")):
            fill = shd.get(qn("w:fill"))
            if fill not in (None, "auto", "FFFFFF", "ffffff"):
                shd.set(qn("w:fill"), "FFFFFF")
        # Coloured borders → black.
        for node in el.iter():
            if node.tag in _MONO_BORDER_TAGS and node.get(qn("w:color")) not in (None, "auto"):
                node.set(qn("w:color"), "000000")


def add_page_break(document: Document) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_fixed_accent_line(
    document: Document,
    *,
    width_cm: float,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    color: str = CYAN,
    size: int = 16,
) -> None:
    """Add a non-editable-looking paragraph border instead of a one-cell table."""
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment
    paragraph.paragraph_format.left_indent = Cm((16.4 - width_cm) / 2 if alignment == WD_ALIGN_PARAGRAPH.CENTER else 0)
    paragraph.paragraph_format.right_indent = Cm((16.4 - width_cm) / 2 if alignment == WD_ALIGN_PARAGRAPH.CENTER else 16.4 - width_cm)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.add_run("\u00a0")
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_front_title(document: Document, title: str, subtitle: str | None = None) -> None:
    paragraph = document.add_paragraph(style="Front Matter Title")
    paragraph.add_run(title)
    if subtitle:
        p = document.add_paragraph(style="Book Subtitle")
        p.add_run(subtitle)
    add_fixed_accent_line(document, width_cm=3.2)


def add_body(document: Document, text: str, *, bold_lead: str | None = None,
             citation: str | None = None, italic: bool = False,
             expand_terms: bool = True,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    if expand_terms:
        text = expand_abbreviations(text)
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.widow_control = True
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, "Cambria", 11.5, TEXT, True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, "Cambria", 11.5, TEXT, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, "Cambria", 11.5, TEXT, italic=italic)
    if citation:
        cite = paragraph.add_run(f" {citation}")
        set_run_font(cite, "Cambria", 11.5, BLUE, True)


def add_bullets(document: Document, items: Iterable[str],
                color: str = CYAN_DARK) -> None:
    for item in items:
        item = expand_abbreviations(item)
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Cm(0.55)
        paragraph.paragraph_format.first_line_indent = Cm(-0.25)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_run_font(run, "Cambria", 11.2, TEXT)


def add_numbered(document: Document, items: Iterable[str]) -> None:
    for item in items:
        item = expand_abbreviations(item)
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Cm(0.65)
        paragraph.paragraph_format.first_line_indent = Cm(-0.3)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_run_font(run, "Cambria", 11.2, TEXT)


def add_info_box(document: Document, title: str, body: str,
                 accent: str = CYAN, icon: str = "KEY POINT") -> None:
    body = expand_abbreviations(body)
    table = document.add_table(1, 2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(2.3)
    table.columns[1].width = Cm(14.1)
    left, right = table.rows[0].cells
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(left, WHITE)
    set_cell_shading(right, WHITE)
    set_cell_margins(left, 160, 120, 160, 120)
    set_cell_margins(right, 160, 180, 160, 180)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(icon)
    set_run_font(r, "Aptos", 9.5, accent, True)
    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = p2.add_run(f"{title}\n")
    set_run_font(title_run, "Aptos", 10.5, accent, True)
    body_run = p2.add_run(body)
    set_run_font(body_run, "Cambria", 10.6, TEXT)
    set_table_borders(table, accent, 8)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def find_chapter_image(filename: str) -> Path | None:
    """Locate a real figure image under chapter-images/."""
    if not CHAPTER_IMAGES.exists():
        return None
    requested = Path(filename)
    direct = CHAPTER_IMAGES / requested
    if direct.exists():
        return direct
    matches = list(CHAPTER_IMAGES.rglob(requested.name))
    return matches[0] if matches else None


def add_real_figure(document: Document, image_path: Path, number: str, caption: str,
                    height_cm: float = 7.1) -> None:
    """Insert an actual figure image (framed, centred) with its caption below."""
    try:
        from PIL import Image
        with Image.open(image_path) as im:
            aspect = im.width / im.height
    except Exception:
        aspect = 1.6
    max_w, max_h = 15.8, 10.6
    if aspect >= 1:
        max_h = min(max_h, max(3.0, height_cm))
    width_cm = max_w
    if width_cm / aspect > max_h:
        width_cm = max_h * aspect

    table = document.add_table(1, 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(width_cm + 0.4)
    cell = table.cell(0, 0)
    cell.width = Cm(width_cm + 0.4)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, WHITE)
    set_cell_margins(cell, 60, 60, 60, 60)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 6, "color": BLACK},
        bottom={"val": "single", "sz": 6, "color": BLACK},
        left={"val": "single", "sz": 6, "color": BLACK},
        right={"val": "single", "sz": 6, "color": BLACK},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))

    caption_p = document.add_paragraph(style="Figure Caption")
    caption_p.add_run(f"{number} — {caption}")


def add_figure_placeholder(document: Document, number: str, caption: str,
                           filename: str, height_cm: float = 7.1,
                           orientation: str = "Landscape 16:9") -> None:
    real = find_chapter_image(filename)
    if real is not None:
        add_real_figure(document, real, number, caption, height_cm)
        return
    table = document.add_table(1, 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.2)
    cell = table.cell(0, 0)
    cell.width = Cm(16.2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, WHITE)
    set_cell_margins(cell, 180, 200, 180, 200)
    row = table.rows[0]
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    set_cell_border(
        cell,
        top={"val": "dashed", "sz": 14, "color": CYAN_DARK},
        bottom={"val": "dashed", "sz": 14, "color": CYAN_DARK},
        left={"val": "dashed", "sz": 14, "color": CYAN_DARK},
        right={"val": "dashed", "sz": 14, "color": CYAN_DARK},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run("IMAGE PLACEHOLDER\n")
    set_run_font(r1, "Aptos Display", 16, CYAN_DARK, True)
    r2 = p.add_run(f"File name: {filename}\n")
    set_run_font(r2, "Consolas", 10, NAVY, True)
    r3 = p.add_run(f"Recommended format: {orientation} · minimum 1920 px")
    set_run_font(r3, "Aptos", 9, MUTED, italic=True)

    caption_p = document.add_paragraph(style="Figure Caption")
    caption_p.add_run(f"{number} — {caption}")


def add_table_caption(document: Document, number: str, caption: str) -> None:
    p = document.add_paragraph(style="Table Caption")
    p.add_run(f"{number} — {caption}")


def add_code_block(document: Document, title: str, code: str) -> None:
    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_before = Pt(4)
    title_paragraph.paragraph_format.space_after = Pt(3)
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, "Aptos", 9.4, CYAN_DARK, True)

    table = document.add_table(1, 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.2)
    cell = table.cell(0, 0)
    set_cell_shading(cell, WHITE)
    set_cell_margins(cell, 120, 140, 120, 140)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 8, "color": BORDER},
        bottom={"val": "single", "sz": 8, "color": BORDER},
        left={"val": "single", "sz": 8, "color": CYAN_DARK},
        right={"val": "single", "sz": 8, "color": BORDER},
    )
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(code.strip())
    set_run_font(run, "Consolas", 7.6, TEXT)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_professional_table(document: Document, headers: Sequence[str],
                           rows: Sequence[Sequence[str]],
                           widths: Sequence[float] | None = None,
                           accent: str = CYAN_DARK,
                           font_size: float = 8.8) -> None:
    effective_font_size = max(font_size, 8.8)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for index, width in enumerate(widths):
            table.columns[index].width = Cm(width)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, WHITE)
        set_cell_margins(cell, 70, 45, 70, 45)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, "Aptos", effective_font_size + 0.5, accent, True)

    for row_index, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for col_index, value in enumerate(row_data):
            cell = row_cells[col_index]
            set_cell_shading(cell, WHITE)
            set_cell_margins(cell, 60, 85, 60, 85)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            expanded_value = re.sub(r"\s*/\s*", " / ", expand_abbreviations(value))
            r = p.add_run(expanded_value)
            set_run_font(r, "Cambria", effective_font_size, TEXT)
    set_table_borders(table, BORDER, 5)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_heading(document: Document, kicker: str, title: str,
                     section_number: str | None = None) -> None:
    p = document.add_paragraph(style="Chapter Kicker")
    p.add_run(kicker.upper())
    heading = document.add_paragraph(style="Heading 2")
    if section_number:
        number_run = heading.add_run(f"{section_number}  ")
        set_run_font(number_run, "Aptos Display", 16, CYAN_DARK, True)
    heading.add_run(title)
    add_fixed_accent_line(
        document,
        width_cm=2.5,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )


def add_chapter_page_break(document: Document, is_last: bool = False) -> None:
    if not is_last:
        add_page_break(document)


def add_cover(document: Document) -> None:
    section = document.sections[0]
    configure_a4(section, 0.65, 0.65, 0.65, 0.65)
    unlink_section(section)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)

    table = document.add_table(1, 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(19.65)
    cell = table.cell(0, 0)
    cell.width = Cm(19.65)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, WHITE)
    set_cell_margins(cell, 360, 540, 360, 540)
    table.rows[0].height = Cm(26.8)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    def ink_run(paragraph, text, size, *, bold=True, soft=False, font="Aptos"):
        run = paragraph.add_run(text)
        set_run_font(run, font, size, INK_MUTED if soft else BLACK, bold)
        return run

    # Original faculty / university logos (kept in their original colours).
    logos = cell.add_table(1, 3)
    logos.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, width in enumerate([Cm(4.2), Cm(7.0), Cm(4.2)]):
        logos.columns[index].width = width
        logos.cell(0, index).width = width
        set_cell_shading(logos.cell(0, index), WHITE)
        logos.cell(0, index).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    lp = logos.cell(0, 0).paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run().add_picture(str(branding_asset("university-logo.jpeg")), width=Cm(2.6))
    rp = logos.cell(0, 2).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp.add_run().add_picture(str(branding_asset("faculty-logo.jpg")), width=Cm(2.6))

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    ink_run(p, FACULTY, 13)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    ink_run(p, f"{UNIVERSITY}, Egypt", 12)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    ink_run(p, "────────", 12)

    # Project emblem.
    emblem = cell.add_paragraph()
    emblem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    emblem.paragraph_format.space_before = Pt(6)
    emblem.paragraph_format.space_after = Pt(2)
    emblem.add_run().add_picture(str(branding_asset("crimelens-logo.png")), width=Cm(3.4))

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    ink_run(p, "CrimeLens", 42, font="Aptos Display")
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    ink_run(p, "Smarter Surveillance, Faster Response", 16, font="Aptos Display")

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    ink_run(
        p,
        "A Graduation Project Submitted in Partial Fulfillment of the Requirements "
        "for the Degree of Bachelor of Science in Computer Science & Artificial Intelligence",
        11,
        soft=True,
    )

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    ink_run(p, "Presented by :", 15, font="Aptos Display")

    names = [member[0] for member in TEAM]
    left_names = names[:4]
    right_names = names[4:]
    name_table = cell.add_table(1, 2)
    name_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    name_table.columns[0].width = Cm(8.8)
    name_table.columns[1].width = Cm(8.8)
    for col, group in enumerate((left_names, right_names)):
        c = name_table.cell(0, col)
        c.width = Cm(8.8)
        set_cell_shading(c, WHITE)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        np = c.paragraphs[0]
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, name in enumerate(group):
            run = np.add_run(name)
            set_run_font(run, "Aptos", 11.5, BLACK, True)
            if i != len(group) - 1:
                np.add_run().add_break()

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    ink_run(p, "Supervised by :", 15, font="Aptos Display")
    for name, role in SUPERVISORS:
        sp = cell.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_after = Pt(0)
        ink_run(sp, f"{role}: {name}", 12)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    ink_run(p, ACADEMIC_YEAR.replace(" / ", " - "), 16, font="Aptos Display")


def add_team_pages(document: Document) -> None:
    add_front_title(document, "Project Team", "Members, identifiers, roles, and primary contributions")
    for index, member in enumerate(TEAM):
        if index == 4:
            add_page_break(document)
            add_front_title(document, "Project Team — Continued")
        table = document.add_table(1, 2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Cm(3.3)
        table.columns[1].width = Cm(12.7)
        photo_cell, detail_cell = table.rows[0].cells
        set_cell_shading(photo_cell, WHITE)
        set_cell_shading(detail_cell, WHITE)
        set_cell_margins(photo_cell, 120, 120, 120, 120)
        set_cell_margins(detail_cell, 140, 180, 140, 180)
        pp = photo_cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.add_run().add_picture(str(team_asset(member[3])), width=Cm(2.45), height=Cm(2.85))
        dp = detail_cell.paragraphs[0]
        dp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name = dp.add_run(member[0] + "\n")
        set_run_font(name, "Aptos Display", 13, NAVY, True)
        sid = dp.add_run(f"University ID: {member[1]}\n")
        set_run_font(sid, "Consolas", 9.5, MUTED)
        role = dp.add_run(f"Role: {member[2]}\n")
        set_run_font(role, "Aptos", 10, CYAN_DARK, True)
        contribution_map = {
            "220357": "Architecture coordination, Laravel integration, web workflows, and end-to-end delivery.",
            "220366": "Flutter officer workflows, mobile API integration, and field interaction states.",
            "220239": "Mobile maps, navigation, officer experience, and service integration.",
            "220051": "Computer-vision models, inference pipeline, model integration, and detection reporting.",
            "220223": "Streaming gateway, MediaMTX/FFmpeg operations, deployment, and service integration.",
            "220031": "Quality assurance, validation scenarios, automated tests, and regression review.",
            "210270": "User-interface architecture, visual language, wireframes, and interaction direction.",
        }
        contribution = dp.add_run("Primary contribution: " + contribution_map[member[1]])
        set_run_font(contribution, "Cambria", 9.4, TEXT)
        set_table_borders(table, BORDER, 6)
        document.add_paragraph().paragraph_format.space_after = Pt(2)
    add_page_break(document)


def add_acknowledgement_page(document: Document) -> None:
    add_front_title(document, "Acknowledgement")
    add_body(
        document,
        "We would "
        "like to express our sincere gratitude to Dr. Mohamed Salah Reda and "
        "Dr. Ahmed El-Naggar for their supervision, technical guidance, constructive "
        "feedback, and continuous encouragement throughout the development of "
        "CrimeLens. We also thank Teaching Assistant Osama Hefny for the time, "
        "coordination, and academic support provided during the project.",
    )
    add_body(
        document,
        "We extend our appreciation to the Faculty of Computers and Artificial "
        "Intelligence at Beni-Suef University for providing the academic environment "
        "in which this project was developed. We are grateful to the teaching staff, "
        "laboratory personnel, colleagues, testers, and every person who contributed "
        "feedback, equipment, sample data, or practical advice.",
    )
    add_body(
        document,
        "Finally, we thank our families for their patience and support during the "
        "months of analysis, implementation, integration, testing, documentation, and "
        "presentation preparation. Their confidence made the demanding stages of this "
        "project easier to complete.",
    )
    add_page_break(document)


def add_abstract_page(document: Document) -> None:
    reset_abbreviation_tracking()
    add_front_title(document, "Abstract")
    add_body(
        document,
        "CrimeLens is an Artificial Intelligence (AI)-assisted Computer-Aided "
        "Dispatch (CAD) and smart-surveillance platform designed to connect video "
        "analysis, police-station operations, field response, and auditable decision "
        "records in one integrated workflow. Conventional Closed-Circuit Television "
        "(CCTV) systems are commonly passive: operators must continuously monitor "
        "large numbers of feeds, while dispatch context, officer availability, and "
        "evidence may remain distributed across separate processes. Fully automated "
        "response is also inappropriate for consequential public-safety decisions "
        "because a detection can be uncertain, context-dependent, or false.",
    )
    add_body(
        document,
        "The project introduces a deliberate Incident layer between machine perception "
        "and operational commitment. Camera streams are relayed through a Python "
        "gateway that supports the Real-Time Streaming Protocol (RTSP), HTTP Live "
        "Streaming (HLS), and Web Real-Time Communication (WebRTC). The AI service "
        "uses a combination of You Only Look Once (YOLO) weapon detection, "
        "Convolutional Neural Network (CNN) classifiers, action classification, "
        "temporal smoothing, and configurable thresholds. Reports are authenticated, "
        "filtered, scored by an explainable six-factor priority engine, and presented "
        "to a human dispatcher through a React/Inertia web console.",
    )
    add_body(
        document,
        "When a dispatcher approves an Incident, the Laravel backend creates an "
        "operational Crime, selects an eligible nearby officer using Redis geospatial "
        "data with PostgreSQL/PostGIS support, and delivers the assignment through "
        "real-time and push-notification channels. A Flutter mobile application "
        "supports field actions, location updates, navigation, panic and Emergency "
        "Distress Signal (SOS) alerts, evidence, "
        "and resolution. Security controls include multi-guard authentication, "
        "Internet Protocol (IP) allow-listing, Hash-Based Message Authentication Code "
        "(HMAC) signatures, "
        "Advanced Encryption Standard (AES)-256 encryption, signed media URLs, "
        "tenant isolation, and an append-only hash-chained decision ledger.",
    )
    add_body(
        document,
        "The resulting architecture demonstrates how AI can improve situational "
        "awareness without becoming the final authority. CrimeLens emphasizes "
        "human-in-the-loop governance, explainability, real-time coordination, "
        "modularity, and accountability across the complete detection-to-response "
        "lifecycle.",
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "Keywords—Artificial Intelligence, Computer-Aided Dispatch, smart surveillance, "
        "video anomaly detection, YOLO, human-in-the-loop, real-time streaming, "
        "geospatial assignment, decision ledger."
    )
    set_run_font(r, "Cambria", 9.8, NAVY, True, italic=True)
    add_page_break(document)


def add_toc_page(document: Document) -> None:
    add_front_title(document, "Table of Contents")

    toc_groups = [
        (
            "PRELIMINARY PAGES",
            [
                ("Project Team", "1", 0),
                ("Acknowledgement", "3", 0),
                ("Abstract", "4", 0),
                ("Table of Contents", "5", 0),
                ("List of Figures", "8", 0),
                ("List of Tables", "12", 0),
            ],
        ),
        (
            "CHAPTER ONE — INTRODUCTION",
            [
                ("Chapter One — Introduction", "1", 0),
                ("1.1 Background", "2", 1),
                ("1.2 Problem Statement", "7", 1),
                ("1.3 Objectives", "9", 1),
                ("1.4 Work Methodology", "13", 1),
                ("1.5 Competition", "16", 1),
                ("1.6 Work Organization", "18", 1),
                ("References", "20", 0),
            ],
        ),
        (
            "CHAPTER TWO — LITERATURE REVIEW",
            [
                ("Chapter Two — Literature Review", "28", 0),
                ("2.1 Background", "29", 1),
                ("2.2 Algorithms", "35", 1),
                ("2.3 Performance Comparison", "49", 1),
                ("2.4 Literature Survey", "57", 1),
                ("Chapter Two References", "58", 0),
            ],
        ),
        (
            "CHAPTER THREE — SYSTEM ANALYSIS",
            [
                ("Chapter Three — System Analysis", "60", 0),
                ("3.1 Introduction to System Analysis", "61", 1),
                ("3.2 Functional Requirements", "69", 1),
                ("3.3 System Models", "84", 1),
                ("Chapter Three References", "120", 0),
            ],
        ),
        (
            "CHAPTER FOUR — SYSTEM DESIGN",
            [
                ("Chapter Four — System Design", "121", 0),
                ("4.1 Introduction to System Design", "122", 1),
                ("4.2 Architecture Overview", "124", 1),
                ("4.3 Module and Interface Design", "128", 1),
                ("4.4 Database and ERD Design", "133", 1),
                ("4.5 Security Design", "141", 1),
                ("4.6 Realtime and UI/UX Design", "145", 1),
                ("4.7 Integration and Operational Design", "148", 1),
                ("Chapter Four References", "153", 0),
            ],
        ),
        (
            "CHAPTER FIVE — SYSTEM IMPLEMENTATION",
            [
                ("Chapter Five — System Implementation", "154", 0),
                ("5.1 Implementation Evidence", "155", 1),
                ("5.2 Laravel Modular Backend", "159", 1),
                ("5.3 Database and Persistence", "161", 1),
                ("5.4 Web Frontend Implementation", "163", 1),
                ("5.5 Incident and Dispatch Implementation", "167", 1),
                ("5.6 Flutter Mobile Implementation", "170", 1),
                ("5.7 AI Model Implementation", "173", 1),
                ("5.8 Gateway and Streaming Implementation", "175", 1),
                ("5.9 Realtime and Queues", "177", 1),
                ("5.10 Security Implementation", "180", 1),
                ("5.11 Runtime and Deployment", "184", 1),
                ("5.12 Testing and Quality", "185", 1),
                ("Chapter Five References", "188", 0),
            ],
        ),
        (
            "CHAPTER SIX — SYSTEM TESTING AND EVALUATION",
            [
                ("Chapter Six — System Testing and Evaluation", "189", 0),
                ("6.1 Evaluation Scope", "190", 1),
                ("6.2 Traceability", "193", 1),
                ("6.3 Functional Validation", "195", 1),
                ("6.4 Access Control", "197", 1),
                ("6.5 Security Evidence", "198", 1),
                ("6.6 Realtime Behavior", "199", 1),
                ("6.7 AI Evaluation", "203", 1),
                ("6.8 Streaming", "205", 1),
                ("6.9 Performance Testing", "206", 1),
                ("6.10 Release Quality", "210", 1),
                ("Chapter Six References", "213", 0),
            ],
        ),
        (
            "CHAPTER SEVEN — CONCLUSION AND FUTURE WORK",
            [
                ("Chapter Seven — Conclusion and Future Work", "214", 0),
                ("7.1 Project Closure", "215", 1),
                ("7.2 Contributions", "217", 1),
                ("7.3 Current Limitations", "219", 1),
                ("7.4 Future Roadmap", "220", 1),
                ("7.5 Ethical and Societal Considerations", "225", 1),
                ("7.6 Final Conclusion", "226", 1),
                ("Chapter Seven References", "227", 0),
            ],
        ),
    ]

    for group_title, entries in toc_groups:
        group = document.add_paragraph()
        group.paragraph_format.space_before = Pt(7)
        group.paragraph_format.space_after = Pt(3)
        run = group.add_run(group_title)
        set_run_font(run, "Aptos", 9.2, CYAN_DARK, True)

        for label, page, level in entries:
            row = document.add_table(1, 2)
            row.alignment = WD_TABLE_ALIGNMENT.CENTER
            row.autofit = False
            row.columns[0].width = Cm(14.7)
            row.columns[1].width = Cm(1.5)
            left, right = row.rows[0].cells
            set_cell_margins(left, 45, 45, 45, 45)
            set_cell_margins(right, 45, 45, 45, 45)

            left_paragraph = left.paragraphs[0]
            left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            left_paragraph.paragraph_format.left_indent = Cm(0.55 * level)
            left_run = left_paragraph.add_run(label)
            set_run_font(
                left_run,
                "Cambria",
                9.5 if level == 0 else 9.2,
                NAVY if level == 0 else TEXT,
                level == 0,
            )

            right_paragraph = right.paragraphs[0]
            right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_run = right_paragraph.add_run(page)
            set_run_font(right_run, "Cambria", 9.4, MUTED, True)

            for cell in (left, right):
                set_cell_border(
                    cell,
                    top={"val": "nil"},
                    bottom={"val": "dotted", "sz": 3, "color": BORDER},
                    left={"val": "nil"},
                    right={"val": "nil"},
                )

    add_page_break(document)


def add_list_page(document: Document, title: str, style_name: str,
                  fallback_items: Sequence[tuple[str, str, int]]) -> None:
    add_front_title(document, title)
    for number, caption, page in fallback_items:
        row = document.add_table(1, 2)
        row.alignment = WD_TABLE_ALIGNMENT.CENTER
        row.autofit = False
        row.columns[0].width = Cm(14.8)
        row.columns[1].width = Cm(1.4)
        left, right = row.rows[0].cells
        lp = left.paragraphs[0]
        lr = lp.add_run(f"{number} — {caption}")
        set_run_font(lr, "Cambria", 9.2, TEXT)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = rp.add_run(str(page))
        set_run_font(rr, "Cambria", 9.2, MUTED, True)
        for cell in (left, right):
            set_cell_border(
                cell,
                top={"val": "nil"}, bottom={"val": "dotted", "sz": 3, "color": BORDER},
                left={"val": "nil"}, right={"val": "nil"},
            )
            set_cell_margins(cell, 40, 40, 40, 40)
    add_page_break(document)


def add_chapter_cover(document: Document, chapter_label: str, title: str,
                      chips: str, description: str, title_size: float = 40) -> None:
    """Render a solid navy brand cover for a chapter opener (no background image).

    The chapter number and title sit on the project's dark identity colour with
    cyan and gold accents, mirroring the book's overall visual language.
    """
    cover_muted = "9FB2C9"
    table = document.add_table(1, 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.4)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, 520, 360, 520, 360)
    table.rows[0].height = Cm(24.0)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    set_cell_border(
        cell,
        top={"val": "single", "sz": 18, "color": CYAN},
        bottom={"val": "single", "sz": 18, "color": GOLD},
        left={"val": "single", "sz": 18, "color": CYAN},
        right={"val": "single", "sz": 18, "color": GOLD},
    )

    # Brand kicker.
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("CRIMELENS")
    set_run_font(r, "Aptos", 13, CYAN, True)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-Assisted CAD and Smart Surveillance")
    set_run_font(r, "Aptos", 10.5, cover_muted)

    # Chapter label and title.
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(chapter_label)
    set_run_font(r, "Aptos", 13, GOLD, True)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    set_run_font(r, "Aptos Display", title_size, LIGHT, True)

    # Gold divider rule.
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("────────")
    set_run_font(r, "Aptos", 12, GOLD, True)

    # Section chips.
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(chips)
    set_run_font(r, "Aptos", 10.5, CYAN)

    # Description.
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    r = p.add_run(description)
    set_run_font(r, "Cambria", 11, cover_muted, italic=True)
    add_page_break(document)


def add_image_chapter_cover(document: Document, image_name: str, chapter_label: str,
                            title: str, chips: str, description: str,
                            title_size: float = 40) -> None:
    """Print-ready (black & white) chapter cover: a themed banner photo at the top
    with the chapter number and title set in black beneath it, on a white page.
    """
    # Framed banner image.
    frame = document.add_table(1, 1)
    frame.alignment = WD_TABLE_ALIGNMENT.CENTER
    frame.autofit = False
    frame.columns[0].width = Cm(16.4)
    fcell = frame.cell(0, 0)
    fcell.width = Cm(16.4)
    set_cell_shading(fcell, WHITE)
    set_cell_margins(fcell, 60, 60, 60, 60)
    set_cell_border(
        fcell,
        top={"val": "single", "sz": 16, "color": BLACK},
        bottom={"val": "single", "sz": 16, "color": BLACK},
        left={"val": "single", "sz": 16, "color": BLACK},
        right={"val": "single", "sz": 16, "color": BLACK},
    )
    ip = fcell.paragraphs[0]
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.space_after = Pt(0)
    ip.add_run().add_picture(str(COVERS / image_name), width=Cm(15.9))

    # Brand kicker.
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("CRIMELENS")
    set_run_font(r, "Aptos", 13, BLACK, True)

    # Chapter label and title.
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(chapter_label)
    set_run_font(r, "Aptos", 14, INK_MUTED, True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_run_font(r, "Aptos Display", title_size, BLACK, True)

    # Divider rule.
    add_fixed_accent_line(document, width_cm=5.0, color=BLACK, size=14)

    # Section chips.
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(chips)
    set_run_font(r, "Aptos", 10.5, BLACK)

    # Description.
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run(description)
    set_run_font(r, "Cambria", 11, INK_MUTED, italic=True)
    add_page_break(document)


def add_chapter_opener(document: Document) -> None:
    add_image_chapter_cover(
        document,
        "chapter1.png",
        "CHAPTER ONE",
        "INTRODUCTION",
        "Background  ·  Problem Statement  ·  Objectives  ·  Methodology  ·  "
        "Competitive Landscape  ·  Work Organization",
        "This chapter establishes the operational, scientific, and engineering "
        "context of CrimeLens and defines the problem the project is designed to solve.",
        title_size=44,
    )


def chapter_page_02(document: Document) -> None:
    add_page_heading(document, "Context and foundations", "Background", "1.1")
    add_body(
        document,
        "Public-safety operations depend on the ability to observe events, understand "
        "their urgency, coordinate limited resources, and preserve an accurate history "
        "of the decisions that were made. Closed-Circuit Television (CCTV) expanded "
        "the observational reach of security organizations, but a traditional camera "
        "installation does not automatically create operational intelligence. It may "
        "record a large amount of video while still depending on a small number of "
        "operators to notice an event, interpret it correctly, identify its location, "
        "contact the appropriate station, and explain the situation to a field officer.",
    )
    add_body(
        document,
        "Artificial Intelligence (AI) and Computer Vision (CV) offer a way to transform "
        "video from passive evidence into an active source of candidate events. Modern "
        "deep-learning methods can extract visual features, detect objects, classify "
        "actions, estimate anomaly scores, and track changes over time. However, the "
        "research literature repeatedly emphasizes that an anomaly is contextual: an "
        "action that is suspicious in one location or time may be normal in another. "
        "Real-world datasets are also difficult to obtain, abnormal events are rare, "
        "and environmental conditions can change the appearance of a scene.",
        citation="[1]",
    )
    add_body(
        document,
        "CrimeLens is positioned at the intersection of smart surveillance and "
        "Computer-Aided Dispatch (CAD). It treats AI as a sensor that reports an "
        "observation, not as an unquestionable authority. The platform connects "
        "camera streams, an AI detection service, an explainable Incident layer, a "
        "human dispatcher, geospatial officer assignment, a mobile field workflow, "
        "and an auditable backend. The result is not simply a detection model or a "
        "camera dashboard; it is an integrated detection-to-response system.",
    )
    add_info_box(
        document,
        "Chapter perspective",
        "The purpose of this background is to explain why a useful surveillance "
        "system must connect perception, human judgment, dispatch operations, field "
        "response, and accountability.",
        CYAN,
    )


def chapter_page_03(document: Document) -> None:
    add_page_heading(document, "From recording to coordinated action",
                     "Evolution of Surveillance and Dispatch")
    add_body(
        document,
        "The earliest digital surveillance workflows primarily improved recording, "
        "storage, and playback. A Video Management System (VMS) enabled operators to "
        "organize cameras, retrieve footage, and investigate events after they were "
        "reported. Later systems added motion detection, object search, camera-health "
        "monitoring, remote viewing, and centralized administration. These capabilities "
        "improved coverage, yet they did not necessarily connect a detected event to "
        "the complete operational lifecycle of a police response.",
    )
    add_body(
        document,
        "A CAD platform addresses a different but complementary problem. It manages "
        "incidents, prioritization, operator ownership, resource assignment, status "
        "updates, and response history. In a conventional process, a camera operator "
        "may observe a suspicious event and communicate it through a manual chain such "
        "as telephone or radio. Each handoff can lose context: the exact location, "
        "camera identity, confidence, available imagery, event age, or nearby officer "
        "status may need to be repeated or reconstructed.",
    )
    add_body(
        document,
        "CrimeLens combines these two directions. Surveillance provides the scene; the "
        "AI layer proposes a candidate event; the CAD layer governs what happens next. "
        "The dispatcher sees the Incident with location, priority factors, camera "
        "context, and available officers. Once approved, the event becomes an "
        "operational Crime and enters the officer lifecycle. This separation preserves "
        "the advantages of early detection without allowing a model output to become "
        "an irreversible operational decision.",
    )
    add_figure_placeholder(
        document,
        "Figure 1.2",
        "Evolution from passive CCTV monitoring to an AI-assisted CAD workflow",
        "CH01_FIG_02_Surveillance_to_CAD_Evolution.png",
        7.0,
    )


def chapter_page_04(document: Document) -> None:
    add_page_heading(document, "Scientific background",
                     "Video Anomaly Detection in Real-World Surveillance")
    add_body(
        document,
        "Video anomaly detection aims to identify events, objects, or behavior that "
        "deviate from an expected pattern. The task is difficult because the word "
        "“anomaly” has no universal visual definition. Running may be expected in a "
        "sports event but alarming inside a restricted facility; a vehicle is normal "
        "on a road but abnormal in a pedestrian area. Therefore, scene context, time, "
        "density, direction, and historical behavior influence interpretation.",
        citation="[1]",
    )
    add_body(
        document,
        "The Artificial Intelligence Review survey supplied with the project groups "
        "the field into areas such as crowd statistics, tracking, and crowd-scene "
        "understanding. It reviews Convolutional Neural Network (CNN) architectures, "
        "Long Short-Term Memory (LSTM) networks, Autoencoder (AE) architectures, "
        "Generative Adversarial Network (GAN) models, U-Net variants, You Only Look "
        "Once (YOLO) detectors, "
        "attention mechanisms, transformers, and hybrid architectures. These families "
        "address different aspects of the problem: spatial appearance, temporal motion, "
        "object localization, reconstruction error, frame prediction, or learned "
        "context.",
        citation="[1]",
    )
    add_body(
        document,
        "Weakly supervised approaches are particularly relevant when long surveillance "
        "videos have only video-level labels. The UCF-Crime work formulated real-world "
        "anomaly detection using Multiple Instance Learning (MIL), where a video is "
        "treated as a bag and temporal segments are treated as instances. A ranking "
        "objective encourages abnormal videos to contain segments with higher anomaly "
        "scores than normal videos, reducing the need for frame-by-frame annotation.",
        citation="[2]",
    )
    add_info_box(
        document,
        "Relationship to CrimeLens",
        "MIL and 3D-CNN approaches are discussed as related research. The current "
        "CrimeLens runtime instead combines YOLO, 2D CNN classifiers, an action "
        "classifier, temporal smoothing, and secured backend reporting.",
        PURPLE,
        "BOUNDARY",
    )


def chapter_page_05(document: Document) -> None:
    add_page_heading(document, "Research taxonomy", "Model Families and Their Roles")
    add_table_caption(document, "Table 1.1", "Major video anomaly-detection model families")
    add_professional_table(
        document,
        ["Family", "Primary strength", "Typical limitation", "CrimeLens relevance"],
        [
            ["2D CNN", "Spatial appearance and frame classification", "Limited explicit temporal modeling", "Crime and fire classifiers"],
            ["3D CNN", "Joint spatial-temporal feature extraction", "Higher data and compute requirements", "Related-work direction"],
            ["LSTM / ConvLSTM", "Temporal sequence modeling", "Training complexity and long-sequence drift", "Alternative temporal model"],
            ["Autoencoder", "Unsupervised learning of normal patterns", "Reconstruction error may be ambiguous", "Research comparison"],
            ["GAN", "Prediction and synthetic generation", "Training instability and weak anomaly explanation", "Research comparison"],
            ["YOLO", "Fast object localization", "Object detection alone does not explain behavior", "Weapon detection"],
            ["Attention / Transformer", "Long-range contextual relationships", "Data and compute intensive", "Future enhancement"],
            ["MIL", "Weak supervision for long videos", "Coarse labels and localization uncertainty", "Literature foundation"],
        ],
        [2.6, 4.2, 4.2, 4.5],
        accent=PURPLE,
        font_size=7.7,
    )
    add_body(
        document,
        "No single model family is universally optimal. A practical system must balance "
        "detection quality, latency, available training data, deployment hardware, "
        "explainability, and the operational consequence of a false alarm. The survey "
        "finds that CNN-based methods remained both popular and competitive during the "
        "reviewed period, while the architectural differences among pre-trained 2D CNN "
        "backbones had an insignificant overall effect in the authors’ statistical "
        "analysis. This supports selecting a model according to deployment constraints "
        "rather than assuming that a deeper backbone is always operationally superior.",
        citation="[1]",
    )


def chapter_page_06(document: Document) -> None:
    add_page_heading(document, "Data and measurement", "Datasets and Evaluation Metrics")
    add_body(
        document,
        "Surveillance anomaly research is shaped by the datasets used for training and "
        "evaluation. The UCF-Crime dataset contains 1,900 long, untrimmed real-world "
        "videos totaling approximately 128 hours and representing 13 categories of "
        "anomalous activity in addition to normal scenes. Other frequently used "
        "benchmarks include UCSD Ped1/Ped2, CUHK Avenue, ShanghaiTech Campus, UMN, "
        "StreetScene, and RWF-2000. Their differences in camera angle, scene count, "
        "crowd density, anomaly type, resolution, and annotation granularity make "
        "direct performance comparisons difficult.",
        citation="[1], [2]",
    )
    add_table_caption(document, "Table 1.2", "Benchmark datasets relevant to surveillance anomaly detection")
    add_professional_table(
        document,
        ["Dataset", "Environment", "Notable characteristic", "Use in the book"],
        [
            ["UCF-Crime", "Real-world, multi-category", "1,900 untrimmed videos; 13 anomaly classes", "Weak supervision and real-world diversity"],
            ["UCSD Ped1/Ped2", "Pedestrian walkways", "Non-pedestrian objects and unusual motion", "Classic anomaly benchmark"],
            ["CUHK Avenue", "Campus avenue", "Camera shake and varied actions", "Frame-level anomaly evaluation"],
            ["ShanghaiTech Campus", "13 scenes", "Lighting, angle, and scene diversity", "Generalization challenge"],
            ["UMN", "Synthetic staged crowd", "Escape/panic behavior", "Crowd anomaly baseline"],
            ["RWF-2000", "Real-world fights", "Binary violent/non-violent clips", "Violence-detection comparison"],
        ],
        [3.0, 3.4, 6.0, 3.2],
        accent=CYAN_DARK,
        font_size=8.0,
    )
    add_body(
        document,
        "Accuracy alone can be misleading for rare-event detection because normal "
        "frames greatly outnumber abnormal frames. Frequently reported measures "
        "include the Area Under the Receiver Operating Characteristic Curve (AUC), "
        "Frame-Level Area Under the Curve "
        "(F-AUC), Pixel-Level Area Under the Curve (P-AUC), Equal Error Rate (EER), "
        "precision, recall, F1 score, false-alarm "
        "rate, and detection delay. The supplied survey identifies F-AUC as a widely "
        "used measure for frame-level localization, while also noting that evaluation "
        "must match the annotation and operational objective.",
        citation="[1]",
    )


def chapter_page_07(document: Document) -> None:
    add_page_heading(document, "Media delivery foundation",
                     "Streaming Protocols in CrimeLens")
    add_body(
        document,
        "A smart-surveillance platform must deliver the same physical camera feed to "
        "consumers with different needs. The AI service requires a stable machine-"
        "readable stream; a dispatcher benefits from low-latency live monitoring; and "
        "mobile or browser clients may require a broadly compatible fallback. "
        "CrimeLens uses a Python gateway based on Flask, the FFmpeg multimedia "
        "framework, and MediaMTX to fan "
        "one camera source into Real-Time Streaming Protocol (RTSP), HTTP Live "
        "Streaming (HLS), and Web Real-Time Communication (WebRTC) delivery formats "
        "instead of opening a separate physical connection for every viewer.",
    )
    add_table_caption(document, "Table 1.3", "Comparison of RTSP, HLS, and WebRTC")
    add_professional_table(
        document,
        ["Protocol", "Expanded name", "Primary role", "Strength", "Trade-off"],
        [
            ["RTSP", "Real-Time Streaming Protocol", "Camera source and AI relay", "Established session control for streaming media", "Browser support is limited"],
            ["HLS", "HTTP Live Streaming", "Compatible web/mobile playback", "HTTP delivery, caching, and broad compatibility", "Segmenting increases latency"],
            ["WebRTC", "Web Real-Time Communication", "Low-latency dispatcher monitoring", "Interactive real-time browser delivery", "Signaling/NAT setup is more complex"],
        ],
        [2.0, 3.8, 3.4, 3.6, 3.2],
        accent=TEAL,
        font_size=8.0,
    )
    add_body(
        document,
        "RTSP Version 2.0 is standardized in RFC 7826, HLS is described in RFC 8216, "
        "and WebRTC is standardized as a browser real-time communication technology. "
        "These protocols are not interchangeable labels: each serves a distinct "
        "transport and client requirement. CrimeLens therefore treats the gateway as "
        "a translation and process-management layer rather than embedding camera "
        "connection logic separately in every application.",
        citation="[3]–[5]",
    )


def chapter_page_08(document: Document) -> None:
    add_page_heading(document, "Operational governance", "Human-in-the-Loop Dispatch")
    add_body(
        document,
        "A detection model produces a probability or label; a dispatch decision commits "
        "people and resources to an operational response. These outputs must not be "
        "treated as equivalent. False positives can arise from unusual lighting, "
        "occlusion, unfamiliar but harmless behavior, dataset bias, or a scene whose "
        "meaning differs from the training context. Conversely, waiting for perfect "
        "certainty can delay the response to a genuine emergency. The system therefore "
        "requires a governance layer that combines machine speed with human authority.",
    )
    add_body(
        document,
        "CrimeLens implements this principle through the Incident entity. Accepted AI "
        "alerts, manual reports, and promoted citizen tips become reviewable Incidents. "
        "Repeated AI alerts from the same camera can reuse a recent pending Incident "
        "instead of flooding the queue. An Incident stores its source, station, "
        "location, confidence, priority score, priority tier, and contributing factors. "
        "A dispatcher claims the Incident, examines the available scene and context, "
        "and then either rejects it as a false alarm or dispatches it. A Crime record "
        "is created only after that commitment.",
    )
    add_body(
        document,
        "The optional auto-dispatch policy is deliberately narrow and is not presented "
        "as the normal workflow. It is intended only for exceptional events that satisfy "
        "all configured gates. This design makes the default system defensible: the AI "
        "accelerates attention and prioritization, while an identifiable operator owns "
        "the decision. Later, the decision ledger can answer who acted, when, on which "
        "Incident, and with which recorded priority factors.",
    )


def chapter_page_09(document: Document) -> None:
    add_page_heading(document, "Project context", "CrimeLens System Overview")
    add_body(
        document,
        "CrimeLens is implemented as a Laravel 13 modular monolith with six bounded "
        "modules: Core, Police, Camera, Gateway, AiModel, and Admin. The Admin and "
        "Station consoles use Inertia.js with React, while field officers use a Flutter "
        "mobile application. PostgreSQL with PostGIS stores durable relational and "
        "spatial data; Redis supports queues, cache, dispatcher presence, and the "
        "high-frequency officer-location path; Laravel Horizon supervises queued work; "
        "Pusher/Laravel Echo updates browser consoles; and Firebase Cloud Messaging "
        "(FCM) provides mobile push delivery.",
    )
    add_body(
        document,
        "The operational workflow begins when the AI service watches gateway RTSP "
        "streams assigned to its machine identity. It can report a suspicious alert "
        "through a signed API request. The backend validates the "
        "identity, IP address, signature, timestamp, nonce, camera assignment, and "
        "detection filter before creating or reusing a pending Incident. The priority "
        "engine calculates a normalized score for new accepted Incidents using "
        "confidence, crime severity, weapon signal, recent area activity, time of day, "
        "and crowd density.",
    )
    add_body(
        document,
        "The Incident appears in real time within the Station web console. Individual "
        "dispatchers can claim, release, dispatch, or reject it under their own "
        "identity. Dispatch creates a Crime and triggers an officer-selection process "
        "using live Redis geospatial data with PostgreSQL/PostGIS as durable spatial "
        "storage and fallback. The officer receives the assignment, navigates to the "
        "scene, updates the lifecycle, communicates with the station, and resolves the "
        "case with notes and evidence where available.",
        citation="[10]",
    )
    add_figure_placeholder(
        document,
        "Figure 1.3",
        "CrimeLens high-level operational workflow",
        "CH01_FIG_03_CrimeLens_High_Level_Workflow.png",
        6.6,
    )


def chapter_page_10(document: Document) -> None:
    add_page_heading(document, "Core challenge definition", "Problem Statement", "1.2")
    add_body(
        document,
        "The central problem addressed by CrimeLens is the absence of a single, "
        "governed workflow that transforms a potential event on a surveillance camera "
        "into an informed, secure, and auditable field response. Traditional camera "
        "systems improve visibility and evidence collection, but they may still rely on "
        "continuous human attention and separate communication channels. Standalone AI "
        "models can detect objects or abnormal behavior, but a model output alone does "
        "not contain the operational authority, resource context, or accountability "
        "required for police dispatch.",
    )
    add_body(
        document,
        "The problem is not only detection accuracy. It includes how to validate a "
        "machine report, prioritize competing events, prevent duplicate handling, "
        "identify an eligible nearby officer, communicate the assignment in real time, "
        "protect sensitive camera and location data, preserve evidence, and record why "
        "a consequential decision was taken. A solution that optimizes only one of "
        "these dimensions may create new risks elsewhere.",
    )
    add_table_caption(document, "Table 1.4", "Problem statement dimensions")
    add_professional_table(
        document,
        ["Dimension", "Current difficulty", "Required system response"],
        [
            ["Perception", "Large video volume and rare abnormal events", "Automated candidate-event detection"],
            ["Context", "A label may not explain urgency or scene conditions", "Location, source, confidence, and priority factors"],
            ["Governance", "Unsafe direct AI-to-response automation", "Human-reviewed Incident layer"],
            ["Coordination", "Manual handoffs lose time and information", "Integrated dispatch and realtime status"],
            ["Resource selection", "Officer availability and proximity change continuously", "Geospatial, eligibility-aware assignment"],
            ["Security", "Machine clients and camera credentials are sensitive", "Layered identity, encryption, signatures, and isolation"],
            ["Accountability", "Decision rationale may be fragmented", "Activity history and tamper-evident ledger"],
        ],
        [3.0, 6.0, 6.2],
        accent=RED,
        font_size=8.1,
    )
    add_info_box(
        document,
        "Formal problem statement",
        "How can an integrated platform use AI-assisted video analysis to reduce the "
        "time and information gap between detection and field response while keeping "
        "dispatch decisions secure, explainable, human-supervised, and auditable?",
        RED,
        "QUESTION",
    )


def chapter_page_11(document: Document) -> None:
    add_page_heading(document, "Human and organizational factors",
                     "Stakeholder Pain Points")
    add_table_caption(document, "Table 1.5", "Stakeholders and operational pain points")
    add_professional_table(
        document,
        ["Stakeholder", "Pain point", "CrimeLens response"],
        [
            ["Dispatcher", "Too many feeds and fragmented incident context", "Priority queue, map, camera context, and review actions"],
            ["Station supervisor", "Limited visibility across operators and resources", "Supervisory web view, health, analytics, and active state"],
            ["Field officer", "Incomplete brief and delayed assignment updates", "Mobile assignment, navigation, status, chat, and evidence"],
            ["Administrator", "Distributed configuration and weak system-health visibility", "Central governance, model/camera health, and reports"],
            ["AI engineer", "Unclear integration contract and unsafe authority boundaries", "Assigned cameras, encrypted payloads, signed report API"],
            ["Citizen", "No structured path for contextual tips", "Public/SMS tip intake with dispatcher triage"],
            ["Auditor / reviewer", "Difficulty reconstructing why a decision occurred", "Priority factors, activity logs, and decision ledger"],
        ],
        [3.0, 6.1, 6.1],
        accent=GOLD,
        font_size=8.0,
    )
    add_body(
        document,
        "These stakeholders do not need the same interface or authority. Administrators "
        "govern the platform but should not handle live incidents. Dispatchers commit "
        "operational decisions from a large-screen web console. Officers need a "
        "field-focused mobile experience with minimal interaction. The AI service "
        "requires a machine API and no user-interface privileges over dispatch or "
        "camera control. This separation reduces accidental authority and keeps each "
        "surface focused on its primary task.",
    )


def chapter_page_12(document: Document) -> None:
    add_page_heading(document, "Operational and technical gaps",
                     "Why Existing Workflows Are Insufficient")
    add_body(
        document,
        "Passive monitoring does not scale linearly. Adding cameras increases the "
        "number of scenes that may contain useful evidence, but it also increases the "
        "attention burden placed on operators. Long periods of normal activity can "
        "reduce vigilance, and abnormal events are statistically rare. Even when an "
        "operator notices an event, the next steps may depend on manual communication "
        "and separate systems for maps, officers, notifications, and reports.",
    )
    add_body(
        document,
        "A standalone anomaly detector is also insufficient. A high confidence score "
        "does not automatically reveal whether an officer is available, whether the "
        "event belongs to the local station, whether the camera is correctly assigned "
        "to the reporting model, whether the same Incident is already being handled, "
        "or whether the media can be accessed securely. The detector must participate "
        "in a larger trust and workflow architecture.",
    )
    add_body(
        document,
        "The final gap concerns the relationship between realtime speed and procedural "
        "integrity. Multiple dispatchers may view the same queue; network delivery may "
        "be delayed; machine clients may retry; and mobile devices may be in the "
        "background. The platform must prevent duplicate dispatch, recover stale "
        "claims, persist important notifications, and keep each tenant isolated even "
        "when WebSocket events are used.",
    )
    add_info_box(
        document,
        "Design implication",
        "CrimeLens must be evaluated as an integrated socio-technical workflow, not "
        "only as an AI classifier or a collection of application screens.",
        RED,
    )


def chapter_page_13(document: Document) -> None:
    add_page_heading(document, "Trust, privacy, and responsibility",
                     "Security and Accountability Problem")
    add_body(
        document,
        "CrimeLens processes camera credentials, live streams, officer locations, "
        "citizen reports, evidence, and operational decisions. A weak security model "
        "could expose surveillance infrastructure, allow a forged AI report, leak data "
        "between stations, or permit an unauthorized user to command a physical camera. "
        "Security must therefore be part of the domain design rather than a final "
        "deployment setting.",
    )
    add_body(
        document,
        "The machine trust problem is especially important. A bearer token alone does "
        "not prove that a request came from the registered AI host or that its body was "
        "not replayed. CrimeLens combines Laravel Sanctum tokens, IP allow-listing, "
        "HMAC-SHA256 request signatures, a timestamp window, and a Redis-cached nonce. "
        "Camera connection values returned to the model use an AES-256-CBC envelope "
        "with integrity verification and a session-specific key. Camera credentials "
        "and model secrets are encrypted at rest.",
    )
    add_body(
        document,
        "Accountability is a related but distinct requirement. Authentication identifies "
        "an actor; an audit trail explains what the actor did and why the system "
        "produced its result. CrimeLens stores priority factors with each Incident, "
        "records activity, and writes consequential decisions to an append-only, "
        "hash-chained ledger. The chain does not make a decision correct, but it makes "
        "silent alteration detectable and supports later review.",
        citation="[10]",
    )


def chapter_page_14(document: Document) -> None:
    add_page_heading(document, "Project goals", "Objectives", "1.3")
    add_body(
        document,
        "The general objective of CrimeLens is to design and implement an integrated "
        "platform that converts surveillance observations into secure, explainable, "
        "human-supervised, and auditable police-response workflows. This objective "
        "requires more than connecting an AI model to a notification endpoint; it "
        "requires coordinated capabilities across cameras, machine identity, incident "
        "governance, dispatch, geospatial services, realtime communication, mobile "
        "field operations, evidence, administration, and testing.",
    )
    add_table_caption(document, "Table 1.6", "General and specific project objectives")
    add_professional_table(
        document,
        ["ID", "Objective", "Expected evidence"],
        [
            ["OBJ-01", "Detect candidate threats from assigned camera streams", "Annotated model output and secured report"],
            ["OBJ-02", "Create a reviewable Incident before operational commitment", "Incident state, source, and dispatcher queue"],
            ["OBJ-03", "Prioritize AI incidents using explainable factors", "Stored score, tier, weights, and factor JSON"],
            ["OBJ-04", "Support safe multi-dispatcher ownership", "Mine/Shared queues and race-safe claim/release"],
            ["OBJ-05", "Assign eligible nearby officers", "Redis GEO/PostGIS-based selection and status checks"],
            ["OBJ-06", "Coordinate mobile field response", "Assignment, navigation, lifecycle, panic, chat, evidence"],
            ["OBJ-07", "Protect sensitive users, services, streams, and data", "Guards, policies, HMAC, encryption, and signed URLs"],
            ["OBJ-08", "Preserve decision accountability", "Activity logs and hash-chained ledger"],
            ["OBJ-09", "Provide analytics and operational visibility", "Dashboards, heatmaps, health, and reports"],
            ["OBJ-10", "Validate critical behavior programmatically", "Laravel, browser, Flutter, and AI-client tests"],
        ],
        [1.8, 8.3, 5.5],
        accent=GOLD,
        font_size=7.8,
    )
    add_info_box(
        document,
        "Measurement principle",
        "An objective is considered demonstrated only by implementation evidence, a "
        "repeatable test, a recorded demo, or a documented benchmark. Design targets "
        "are not reported as achieved results without measurement.",
        GREEN,
    )


def chapter_page_15(document: Document) -> None:
    add_page_heading(document, "Detailed goals", "Specific Objectives and Success Criteria")
    add_numbered(
        document,
        [
            "Provide separate, role-appropriate surfaces for administrators, station dispatchers, and field officers.",
            "Ensure the AI service can access only its assigned cameras and can only report observations.",
            "Keep the dispatcher as the normal authority for approving or rejecting an AI-originated Incident.",
            "Store a transparent priority explanation rather than an unexplained urgency label.",
            "Prevent duplicate claims and duplicate Crime creation under concurrent operator actions.",
            "Use realtime browser events and mobile push while preserving persistent application state.",
            "Support RTSP, HLS, and WebRTC through a managed gateway instead of duplicating camera connections.",
            "Maintain station isolation in HTTP queries, policies, and private broadcast channels.",
            "Provide health visibility for models, cameras, gateway processes, queues, and scheduled work.",
            "Document known integration gaps and future work rather than presenting them as completed behavior.",
        ],
    )
    add_body(
        document,
        "Success criteria span functional and non-functional evidence. Functional "
        "criteria include the ability to ingest an AI report, create and score an "
        "Incident, claim it, dispatch or reject it, assign an officer, deliver the "
        "assignment, and update the case. Non-functional criteria include authorization, "
        "transaction safety, tenant isolation, stream reuse, queue observability, "
        "low-latency geospatial lookup, maintainable module boundaries, and testable "
        "business rules.",
    )
    add_figure_placeholder(
        document,
        "Figure 1.4",
        "CrimeLens objective hierarchy",
        "CH01_FIG_04_Objective_Hierarchy.png",
        6.6,
    )
    add_info_box(
        document,
        "Expected impact",
        "The project aims to reduce information loss and operator fragmentation while "
        "increasing situational awareness, accountability, and coordination—not to "
        "remove human responsibility.",
        GOLD,
    )


def chapter_page_16(document: Document) -> None:
    add_page_heading(document, "System boundary", "Project Scope")
    add_body(
        document,
        "Scope defines what the graduation project demonstrates and prevents future "
        "possibilities from being confused with current behavior. CrimeLens covers the "
        "software path from camera and AI integration through dispatcher review and "
        "field response. It is an academic prototype with a broad integrated feature "
        "set, not a certified national emergency service or a replacement for every "
        "existing police information system.",
    )
    add_table_caption(document, "Table 1.7", "In-scope and out-of-scope boundaries")
    add_professional_table(
        document,
        ["In scope", "Out of scope / future validation"],
        [
            ["Camera registration, streaming gateway, HLS/WebRTC viewing, selected PTZ/alarm operations", "Certification for every camera manufacturer and public-network deployment"],
            ["AI machine identity, assigned cameras, heartbeat, signed alert reporting", "An AI model with legal authority to dispatch or command devices directly"],
            ["Incident intake, priority scoring, claim/release, dispatch/reject, manual incidents", "A fully autonomous police-response system"],
            ["Nearest-officer selection, mobile assignment lifecycle, GPS, panic, chat, evidence", "Integration with national identity, court, or emergency-number infrastructure"],
            ["Admin/station analytics, health, reports, activity logs, decision ledger", "Production-scale nationwide benchmark and disaster-recovery certification"],
            ["Pest, browser, Flutter, and isolated AI-client tests", "A claim that every discovered test passed without a recorded full environment run"],
        ],
        [7.8, 7.8],
        accent=BLUE,
        font_size=8.1,
    )
    add_figure_placeholder(
        document,
        "Figure 1.5",
        "Project scope boundary",
        "CH01_FIG_05_Project_Scope_Boundary.png",
        5.0,
    )


def chapter_page_17(document: Document) -> None:
    add_page_heading(document, "Value of the work", "Expected Contributions and Impact")
    add_body(
        document,
        "The first contribution is architectural: CrimeLens introduces an explicit "
        "Incident governance layer between AI detection and Crime creation. This "
        "boundary allows the system to preserve machine speed while keeping operational "
        "authority with a dispatcher. It also creates a consistent entry point for AI "
        "reports, manual reports, and promoted citizen tips.",
    )
    add_body(
        document,
        "The second contribution is explainability. The priority engine combines six "
        "normalized factors—confidence, crime-type severity, weapon signal, repeat-area "
        "activity, time of day, and crowd density—and stores the complete breakdown. "
        "The third contribution is operational integration: the project connects "
        "streaming, realtime web interfaces, geospatial officer selection, mobile field "
        "response, and persistent records across one platform.",
    )
    add_body(
        document,
        "The fourth contribution is accountability and security. AI and gateway clients "
        "operate behind explicit trust boundaries; data is tenant-scoped; sensitive "
        "media uses signed access; and consequential decisions are recorded in a "
        "tamper-evident ledger. Finally, the project contributes an extensive modular "
        "test inventory and a documented distinction between implemented behavior, "
        "acceptance targets, and future enhancements.",
    )
    add_info_box(
        document,
        "Academic contribution",
        "CrimeLens demonstrates how research ideas in anomaly detection can be placed "
        "inside a complete software-engineering and operational-governance context.",
        PURPLE,
    )


def chapter_page_18(document: Document) -> None:
    add_page_heading(document, "Development strategy", "Work Methodology", "1.4")
    add_body(
        document,
        "The project followed an Agile, iterative, integration-oriented methodology. "
        "The team first analyzed the problem domain, identified the human and machine "
        "actors, and separated perception from operational authority. System models "
        "and module boundaries were then created around incidents, dispatch, field "
        "operations, cameras, AI integration, administration, communication, and audit. "
        "Features were implemented as vertical slices so that each increment could "
        "connect a user interface to backend rules, persistence, realtime events, "
        "tests, and demo evidence.",
    )
    add_body(
        document,
        "This methodology was not a one-time linear sequence. Implementation evidence "
        "caused the design and documentation to evolve. For example, the dispatcher "
        "surface moved to an Inertia/React web console because multi-panel monitoring "
        "and realtime interaction are better suited to a large screen. The Incident "
        "layer was strengthened to prevent direct AI-to-Crime assignment. Security, "
        "authorization, concurrency, and audit were treated as cross-cutting concerns "
        "throughout the lifecycle.",
    )
    add_figure_placeholder(
        document,
        "Figure 1.6",
        "Iterative development methodology",
        "CH01_FIG_06_Iterative_Methodology_Cycle.png",
        8.0,
    )
    add_info_box(
        document,
        "Methodology principle",
        "Design, implementation, tests, and documentation must describe the same "
        "system. When evidence changes, the model and documentation are updated.",
        GREEN,
    )


def chapter_page_19(document: Document) -> None:
    add_page_heading(document, "Phases and deliverables", "Methodology Stages")
    add_table_caption(document, "Table 1.8", "Methodology phases and outputs")
    add_professional_table(
        document,
        ["Phase", "Main activities", "Primary outputs"],
        [
            ["1. Discovery", "Problem analysis, stakeholder identification, literature review, technology constraints", "Problem statement, actors, scope, risks"],
            ["2. Requirements", "Functional and non-functional requirement definition", "SRS, authority matrix, acceptance targets"],
            ["3. Modeling", "Workflow, states, data model, architecture, security boundaries", "Use cases, DFD/UML, ERD, module map"],
            ["4. Foundation", "Modules, database, authentication, shared conventions", "Laravel core, schemas, guards, base interfaces"],
            ["5. Sensing", "Gateway, camera control, AI identity, encrypted streams", "RTSP/HLS/WebRTC pipeline and model API"],
            ["6. Dispatch", "Incident layer, priority engine, console, realtime", "Claim/review/dispatch/reject workflow"],
            ["7. Field response", "Officer mobile, GPS, panic, chat, evidence", "Complete officer lifecycle"],
            ["8. Intelligence", "Analytics, patterns, ledger, reports, health", "Operational insights and accountability"],
            ["9. Verification", "Feature, unit, browser, mobile, and integration tests", "Test evidence and defect correction"],
            ["10. Presentation", "Seeded data, screenshots, documentation, demo rehearsal", "Book, presentation, and live demo"],
        ],
        [2.6, 7.0, 6.0],
        accent=GREEN,
        font_size=7.8,
    )
    add_body(
        document,
        "Each phase produced a usable artifact rather than an isolated document. "
        "Requirements were linked to routes, services, interfaces, and tests. "
        "Architecture decisions were validated against the actual repository. Known "
        "limitations—such as the incomplete persistence link for AI event windows in "
        "the newer Incident path and the unscheduled stale-Crime escalation command—"
        "are documented explicitly rather than omitted.",
    )
    add_info_box(
        document,
        "Team coordination",
        "The seven workstreams used shared API contracts, event payloads, naming "
        "conventions, demo data, and integration checkpoints to prevent the backend, "
        "web, mobile, AI, streaming, testing, and design work from becoming separate projects.",
        CYAN,
    )


def chapter_page_20(document: Document) -> None:
    add_page_heading(document, "Project execution", "Work Plan and Team Coordination")
    add_body(
        document,
        "The historical execution plan organized the project into five monthly "
        "increments from January to May 2026. January established the modular backend, "
        "database, authentication, and client foundations. February focused on cameras, "
        "streaming, AI identity, encryption, and model reporting. March delivered the "
        "Incident layer, priority engine, dispatcher console, and realtime backbone. "
        "April closed the field-response loop through the officer application, "
        "geospatial assignment, panic/BOLO, citizen tips, chat, and evidence. May "
        "concentrated on analytics, the decision ledger, reporting, testing, demo data, "
        "documentation, and visual polish.",
    )
    add_body(
        document,
        "The dependency structure was intentional. The officer application could not "
        "complete assignment flows until the Incident and Crime contracts were stable; "
        "the dispatcher console depended on realtime channels and role-specific "
        "authentication; and the AI service required camera assignment, gateway URLs, "
        "encryption, and backend validation. Integration checkpoints were therefore "
        "scheduled at the end of each major increment.",
    )
    add_figure_placeholder(
        document,
        "Figure 1.7",
        "Five-month work plan",
        "CH01_FIG_07_Five_Month_Timeline.png",
        8.2,
    )
    add_bullets(
        document,
        [
            "P0 critical path: authentication, data model, streaming, AI integration, Incident layer, dispatcher console, officer assignment.",
            "P1 operational support: citizen tips, escalation, panic/SOS, BOLO, and chat.",
            "P2/P3 differentiation: heatmaps, analytics, pattern detection, tamper monitoring, ledger, reports, and simulation.",
        ],
    )


def chapter_page_21(document: Document) -> None:
    add_page_heading(document, "Verification approach", "Testing and Validation Methodology")
    add_body(
        document,
        "Testing is treated as part of the feature definition. The repository contains "
        "a feature-heavy Pest suite that exercises requests through routing, middleware, "
        "controllers, services, and persistence, while unit tests cover pure logic such "
        "as geometry, scoring, and ledger integrity. Browser tests exercise the React "
        "consoles, Flutter tests cover mobile behavior, and isolated Python tests verify "
        "parts of the AI client and model integration.",
        citation="[10]",
    )
    add_body(
        document,
        "Critical scenarios include guard-specific login, forced password rotation, "
        "station isolation, AI IP and signature validation, encrypted camera payloads, "
        "Incident creation, race-safe claim/release, single-Crime dispatch guarantees, "
        "officer assignment, panic, camera operations, signed media, realtime payloads, "
        "and decision-ledger immutability. The production web build and formatting "
        "checks form additional release gates.",
    )
    add_body(
        document,
        "Performance values are handled differently from functional tests. Targets such "
        "as camera-to-notification time, nearest-officer lookup latency, stream startup, "
        "and normal API response time require a recorded environment, dataset size, "
        "network condition, and measurement method. They remain acceptance targets "
        "until measured, which keeps architectural intention clearly separated from "
        "measured experimental evidence.",
    )
    add_info_box(
        document,
        "Current repository inventory",
        "The current backend evidence records 883 passing Pest cases across 148 "
        "Laravel test files, while the mobile repository contains 38 Flutter test "
        "files. Inventory counts are reported separately from executed-suite results.",
        GREEN,
        "EVIDENCE",
    )
    add_bullets(
        document,
        [
            "Functional validation: Does the workflow produce the correct state and authorization result?",
            "Integration validation: Do services, clients, queues, streams, and events cooperate correctly?",
            "Operational validation: Does the demo remain understandable and recoverable under realistic conditions?",
            "Research validation: Are AI metrics derived from the project’s own experiments rather than copied from related work?",
        ],
    )


def chapter_page_22(document: Document) -> None:
    add_page_heading(document, "External context",
                     "Competition with Other Programs", "1.5")
    add_body(
        document,
        "CrimeLens occupies a space shared by several established product categories "
        "rather than one directly identical competitor. Video Management Systems focus "
        "on camera administration, recording, playback, and surveillance operations. "
        "Unified security platforms combine video with access control and other "
        "physical-security functions. Public-safety command-center and CAD platforms "
        "coordinate calls, incidents, responders, and situational information. "
        "Cloud-managed video platforms emphasize centralized deployment, remote access, "
        "and analytics.",
    )
    add_body(
        document,
        "Representative enterprise vendors include Genetec, Motorola Solutions, Axon, "
        "and Verkada. Their products differ in scope, maturity, deployment model, "
        "ecosystem, and intended customer. The purpose of the comparison is therefore "
        "not to claim feature superiority over mature commercial systems. It is to "
        "identify the specific academic contribution of CrimeLens: an integrated, "
        "explainable, human-supervised path from AI-originated camera evidence to field "
        "response and an auditable decision history.",
        citation="[6]–[9]",
    )
    add_figure_placeholder(
        document,
        "Figure 1.8",
        "Market-category positioning map",
        "CH01_FIG_08_Market_Category_Positioning.png",
        8.4,
    )
    add_info_box(
        document,
        "Basis of comparison",
        "The comparison considers product emphasis and architectural positioning rather "
        "than asserting that any enterprise platform lacks a specific capability, since "
        "such an assertion would depend on the exact edition and the official "
        "documentation of each product.",
        PURPLE,
    )


def chapter_page_23(document: Document) -> None:
    add_page_heading(document, "Comparative evaluation",
                     "Solution Categories and CrimeLens Position")
    add_table_caption(document, "Table 1.9", "High-level comparison of solution categories")
    add_professional_table(
        document,
        ["Criterion", "Traditional CCTV / VMS", "Enterprise unified platform", "CrimeLens academic focus"],
        [
            ["Primary emphasis", "Observation, recording, playback", "Broad security or public-safety ecosystem", "Detection-to-response workflow"],
            ["AI intake", "Optional or external analytics", "Product/integration dependent", "Secured AI report → Incident"],
            ["Human review", "Usually external manual process", "Depends on operational product", "Explicit dispatcher authority"],
            ["Priority explanation", "Not usually a core CCTV function", "Implementation dependent", "Six-factor stored breakdown"],
            ["Field workflow", "Outside the camera platform", "Dedicated modules/integrations", "Flutter officer lifecycle"],
            ["Streaming", "Vendor VMS/client delivery", "Platform-specific", "RTSP relay + HLS + WebRTC"],
            ["Geospatial assignment", "Outside typical VMS scope", "CAD/public-safety dependent", "Redis GEO + PostGIS"],
            ["Audit", "Video and operator logs", "Enterprise audit capabilities", "Activity log + hash chain"],
            ["Deployment goal", "Local to enterprise surveillance", "Enterprise/city scale", "Research prototype and graduation demo"],
        ],
        [3.0, 4.2, 4.5, 4.2],
        accent=PURPLE,
        font_size=7.5,
    )
    add_body(
        document,
        "The comparison highlights a design gap that the project explores: many camera "
        "products are strongest at sensing and evidence, while CAD products are "
        "strongest at incident coordination. CrimeLens connects these concerns and "
        "places an explainable Incident boundary between them. Its modular-monolith "
        "architecture and open technology stack make the complete workflow visible for "
        "academic analysis, testing, and demonstration.",
    )
    add_body(
        document,
        "Commercial maturity, certifications, support, integrations, hardware coverage, "
        "and nationwide scalability are outside the graduation prototype’s present "
        "claim. The project's value lies in its coherent implementation, documented "
        "authority model, security layers, testability, and the ability to explain how "
        "a detection becomes—or does not become—an operational response.",
    )


def chapter_page_24(document: Document) -> None:
    add_page_heading(document, "Differentiation", "CrimeLens Value Proposition")
    add_body(
        document,
        "CrimeLens is summarized by the principle: “AI reports; humans decide; the "
        "backend acts.” This principle creates a clear separation of responsibilities. "
        "The AI service is useful but untrusted; the dispatcher has decision authority; "
        "the backend enforces transactions, tenancy, permissions, and camera commands; "
        "the officer executes the field workflow; and administrators govern the "
        "platform without silently becoming dispatchers.",
    )
    add_body(
        document,
        "The most important differentiators are not isolated user-interface features. "
        "They are the Incident governance layer, explainable priority factors, "
        "race-safe dispatcher ownership, geospatial assignment, backend-mediated "
        "physical-device authority, secured machine integration, realtime coordination, "
        "and the tamper-evident decision ledger. Together, these mechanisms address "
        "speed and accountability as complementary requirements.",
    )
    add_bullets(
        document,
        [
            "Proactive: camera and AI services can surface candidate events before a manual call.",
            "Human-supervised: normal dispatch remains an explicit operator decision.",
            "Explainable: priority is decomposed into stored factors and weights.",
            "Connected: web consoles, mobile response, streaming, queues, maps, and notifications share one domain workflow.",
            "Auditable: important state transitions and decisions remain reconstructable.",
            "Honest: known gaps and unmeasured targets are documented as future work.",
        ],
    )


def chapter_page_25(document: Document) -> None:
    add_page_heading(document, "Document roadmap", "Work Organization", "1.6")
    add_body(
        document,
        "The remainder of this graduation book is organized to move from scientific "
        "context to system evidence. Each chapter has a distinct purpose and should "
        "avoid repeating material unless a short summary is needed to connect the "
        "argument. Figures, tables, abbreviations, requirements, diagrams, interfaces, "
        "test results, and references will follow the numbering and citation system "
        "established in this first chapter.",
    )
    add_table_caption(document, "Table 1.10", "Book chapters and their purpose")
    add_professional_table(
        document,
        ["Chapter", "Title", "Primary content"],
        [
            ["One", "Introduction", "Background, problem, objectives, methodology, competition, and book organization"],
            ["Two", "Literature Review", "Video anomaly detection, datasets, model families, metrics, and related work"],
            ["Three", "System Analysis", "Actors, requirements, use cases, workflows, UML/DFD, constraints, and data gathering"],
            ["Four", "System Design", "Architecture, security, module design, ERD/schema, streaming, UI/UX, and interfaces"],
            ["Five", "System Implementation", "Laravel, React/Inertia, Flutter, AI service, gateway, database, realtime, and deployment"],
            ["Six", "System Testing and Evaluation", "Test strategy, functional validation, AI experiments, performance measurements, and results"],
            ["Seven", "Conclusion and Future Work", "Achievements, limitations, roadmap, scalability, research directions, and final conclusions"],
            ["Appendices", "Supporting Material", "API samples, requirement matrices, setup, screenshots, selected code, and additional results"],
        ],
        [2.3, 4.0, 9.4],
        accent=CYAN_DARK,
        font_size=8.1,
    )
    add_body(
        document,
        "Chapter Two will use the supplied 2025 survey as a central reference while "
        "also separating general anomaly-detection research, UCF-Crime/MIL, 3D-CNN "
        "approaches, and the models actually integrated into CrimeLens. Experimental "
        "claims will be supported by the project’s own notebooks, logs, and recorded "
        "evaluation rather than borrowed performance values.",
        citation="[1], [2]",
    )


def add_references(document: Document) -> None:
    add_page_break(document)
    p = document.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("References")
    for index, (number, reference) in enumerate(REFERENCES):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.first_line_indent = Cm(-0.8)
        paragraph.paragraph_format.space_after = Pt(6)
        n = paragraph.add_run(number + " ")
        set_run_font(n, "Cambria", 9.5, NAVY, True)
        r = paragraph.add_run(reference)
        set_run_font(r, "Cambria", 9.5, TEXT)


def add_chapter_two_opener(document: Document) -> None:
    add_image_chapter_cover(
        document,
        "chapter2.png",
        "CHAPTER TWO",
        "LITERATURE REVIEW",
        "Background  ·  Algorithms  ·  Datasets  ·  Evaluation Metrics  ·  "
        "Performance Comparison  ·  Literature Synthesis",
        "This chapter reviews the scientific foundations of automated surveillance-video "
        "analysis and identifies the methods, evidence, limitations, and research gaps "
        "most relevant to crime and anomaly detection.",
        title_size=38,
    )


def chapter_two_page_specs() -> list[dict]:
    return [
        {
            "kicker": "Review scope",
            "title": "Background and Scope of the Literature Review",
            "section": "2.1",
            "paragraphs": [
                (
                    "The literature surrounding automated crime and anomaly detection "
                    "covers several related but non-identical tasks. Video anomaly "
                    "detection attempts to identify observations that depart from an "
                    "expected pattern; action recognition assigns a known activity label; "
                    "object detection localizes entities such as people, vehicles, fire, "
                    "or weapons; violence detection concentrates on aggressive physical "
                    "behavior; and crowd analysis estimates density, motion, flow, and "
                    "collective behavior. A reliable review must preserve these "
                    "distinctions because a method that recognizes a weapon is not "
                    "automatically capable of understanding intent, and a model that "
                    "detects unusual motion does not necessarily identify a legal crime.",
                    "[1]",
                ),
                (
                    "This chapter therefore examines the scientific methods that convert "
                    "video into evidence: spatial and temporal feature extraction, "
                    "convolutional and recurrent networks, reconstruction and prediction "
                    "models, weak supervision, object detection, attention, and "
                    "transformers. It also reviews datasets and evaluation "
                    "metrics because reported performance is meaningful only when the "
                    "task definition, annotation level, test protocol, and operating "
                    "threshold are known. System requirements, actors, use cases, database "
                    "design, software architecture, interfaces, and implementation details "
                    "are intentionally reserved for later chapters.",
                    None,
                ),
            ],
            "table": {
                "number": "Table 2.1",
                "caption": "Scope and boundaries of the literature review",
                "headers": ["Included in Chapter Two", "Deferred to later chapters"],
                "rows": [
                    ["Anomaly, violence, action, and object-detection research", "CrimeLens functional and non-functional requirements"],
                    ["Learning paradigms, model families, losses, and features", "Actors, use cases, sequence flows, and system models"],
                    ["Datasets, annotation schemes, metrics, and reported results", "Software architecture, database schema, and interfaces"],
                    ["Research limitations, generalization, and open problems", "Implementation, deployment, project testing, and measured results"],
                ],
                "widths": [8.1, 8.1],
                "accent": CYAN_DARK,
                "font_size": 8.4,
            },
        },
        {
            "kicker": "Historical development",
            "title": "From Handcrafted Features to Deep Video Models",
            "paragraphs": [
                (
                    "Early surveillance-analysis systems relied on manually designed "
                    "descriptors of appearance and motion. Background subtraction separated "
                    "moving foreground regions; optical flow represented motion direction "
                    "and magnitude; trajectories described object movement; and histograms "
                    "summarized gradients, flow orientations, or local spatiotemporal "
                    "patterns. These methods were computationally understandable and could "
                    "work in constrained scenes, but their assumptions were fragile under "
                    "camera motion, illumination changes, occlusion, weather, compression "
                    "noise, and complex interactions among multiple people.",
                    "[1]",
                ),
                (
                    "Deep learning changed the emphasis from manual feature engineering to "
                    "representation learning. Two-dimensional convolutional networks learned "
                    "appearance from frames, recurrent models summarized temporal sequences, "
                    "and three-dimensional convolution learned joint spatial-temporal "
                    "filters. Autoencoders and future-frame predictors modeled normality "
                    "without requiring examples of every anomaly, while multiple-instance "
                    "learning made it possible to train on long videos with only video-level "
                    "labels. More recent attention and transformer architectures model "
                    "long-range relationships among objects, patches, and temporal segments. "
                    "The progression is not a simple replacement sequence: practical systems "
                    "still combine classical motion cues, convolutional features, temporal "
                    "models, object detectors, and decision rules.",
                    "[11]–[15], [20], [21]",
                ),
            ],
            "figure": (
                "Figure 2.2",
                "Evolution of surveillance-video analysis methods",
                "CH02_FIG_02_Method_Evolution_Timeline.png",
                7.1,
            ),
        },
        {
            "kicker": "Problem definition",
            "title": "What Constitutes an Anomaly?",
            "paragraphs": [
                (
                    "An anomaly is not defined only by visual rarity. It may be a rare "
                    "object, an unusual action, an unexpected trajectory, an abnormal crowd "
                    "density, a violation of scene rules, or a familiar action occurring at "
                    "an inappropriate time or place. Running can be normal in a sports area "
                    "and suspicious inside a restricted facility. A vehicle is expected on "
                    "a roadway but anomalous on a pedestrian path. Context therefore acts as "
                    "part of the label, even when it is absent from the raw class name.",
                    "[1]",
                ),
                (
                    "The literature also distinguishes point anomalies, contextual "
                    "anomalies, and collective anomalies. A point anomaly can be identified "
                    "from an individual frame or object; a contextual anomaly is abnormal "
                    "only under specific spatial or temporal conditions; and a collective "
                    "anomaly emerges from a sequence or group even when individual elements "
                    "appear ordinary. This distinction affects model design. Frame "
                    "classifiers may capture strong point evidence, whereas trajectory, "
                    "recurrent, three-dimensional convolutional, or attention-based methods "
                    "are better suited to relationships that unfold over time.",
                    None,
                ),
            ],
            "table": {
                "number": "Table 2.2",
                "caption": "Categories of video anomalies",
                "headers": ["Category", "Description", "Illustrative example", "Modeling need"],
                "rows": [
                    ["Point", "A single observation differs strongly from normal data", "Visible weapon or fire", "Spatial detection/classification"],
                    ["Contextual", "An event is abnormal only in its scene or time", "Vehicle in a pedestrian zone", "Scene and metadata context"],
                    ["Collective", "A sequence or group pattern becomes abnormal", "Panic, fight, or crowd dispersal", "Temporal/group modeling"],
                    ["Object-level", "An entity, state, or interaction is suspicious", "Abandoned object or forced entry", "Detection, tracking, relations"],
                ],
                "widths": [2.5, 5.0, 4.2, 4.4],
                "accent": PURPLE,
                "font_size": 7.8,
            },
            "figure": (
                "Figure 2.3",
                "Contextual taxonomy of video anomalies",
                "CH02_FIG_03_Anomaly_Context_Taxonomy.png",
                3.7,
            ),
        },
        {
            "kicker": "Learning settings",
            "title": "Supervision Paradigms",
            "paragraphs": [
                (
                    "Fully supervised learning uses explicit class or temporal labels and "
                    "can directly optimize discrimination, but detailed video annotation is "
                    "expensive and ambiguous. Unsupervised and one-class methods learn a "
                    "representation of normal behavior and treat large deviations as "
                    "anomalies. Semi-supervised anomaly detection commonly trains only on "
                    "normal clips while evaluating on mixed normal and abnormal scenes. "
                    "Weakly supervised methods use coarse video-level labels and infer which "
                    "segments caused the label. Self-supervised methods construct proxy tasks "
                    "such as temporal ordering, transformation prediction, masked modeling, "
                    "or future-frame prediction.",
                    "[1], [2], [11], [12]",
                ),
                (
                    "The choice of supervision controls both data cost and semantic "
                    "specificity. Normality models can discover unseen deviations but may "
                    "flag harmless novelty. Fully supervised classifiers can distinguish "
                    "known categories but often fail on unrepresented events. Weak "
                    "supervision provides a practical compromise for long surveillance "
                    "videos, although localization remains uncertain because the model is "
                    "not told where the event occurs. Consequently, the best paradigm "
                    "depends on whether the objective is open-set anomaly discovery, known "
                    "crime classification, temporal localization, or real-time alerting.",
                    None,
                ),
            ],
            "table": {
                "number": "Table 2.3",
                "caption": "Comparison of supervision paradigms",
                "headers": ["Paradigm", "Training labels", "Main advantage", "Principal risk"],
                "rows": [
                    ["Fully supervised", "Frame, clip, box, or class labels", "Direct task optimization", "High annotation cost; closed-set bias"],
                    ["Semi-supervised / one-class", "Primarily normal training data", "No need to enumerate every anomaly", "Normal novelty may trigger alarms"],
                    ["Weakly supervised", "Video-level normal/abnormal labels", "Scales to long untrimmed video", "Imprecise temporal localization"],
                    ["Self-supervised", "Labels generated from proxy tasks", "Uses large unlabeled collections", "Proxy objective may not match operations"],
                ],
                "widths": [3.1, 4.1, 4.5, 4.5],
                "accent": TEAL,
                "font_size": 7.9,
            },
            "figure": (
                "Figure 2.4",
                "Supervision paradigms for video anomaly detection",
                "CH02_FIG_04_Supervision_Paradigms.png",
                3.8,
            ),
        },
        {
            "kicker": "Input representation",
            "title": "Feature Representations for Surveillance Video",
            "paragraphs": [
                (
                    "A model does not consume the abstract idea of a crime; it consumes a "
                    "representation. Appearance features encode color, texture, shape, and "
                    "objects from individual frames. Motion features use frame differences, "
                    "optical flow, trajectories, or learned temporal filters. Object-centric "
                    "representations describe detected people and items, their tracks, "
                    "poses, distances, and interactions. Scene representations encode zones, "
                    "entry directions, density, time, and environmental context.",
                    "[1], [15]",
                ),
                (
                    "Representation quality determines what a later classifier can learn. "
                    "Frame features are efficient but may confuse visually similar actions. "
                    "Optical flow emphasizes motion while suppressing appearance, yet it is "
                    "sensitive to camera movement and estimation cost. Object-centric "
                    "features are explainable but depend on detector and tracker reliability. "
                    "Spatiotemporal tensors preserve richer evidence but increase memory and "
                    "computation.",
                    None,
                ),
            ],
            "table": {
                "number": "Table 2.4",
                "caption": "Feature representations used in surveillance analysis",
                "headers": ["Representation", "Captures", "Typical method", "Limitation"],
                "rows": [
                    ["RGB appearance", "Objects, texture, scene", "2D CNN / detector", "Weak temporal evidence"],
                    ["Motion", "Direction and movement", "Optical flow / temporal difference", "Noise and camera-motion sensitivity"],
                    ["Spatiotemporal volume", "Joint appearance and motion", "3D CNN / video transformer", "Compute and data demand"],
                    ["Object graph", "Entities and interactions", "Detection, tracking, relation model", "Error propagation"],
                ],
                "widths": [3.0, 3.7, 4.5, 5.0],
                "accent": CYAN_DARK,
                "font_size": 7.7,
            },
            "figure": (
                "Figure 2.5",
                "Spatial, temporal, and object feature representations",
                "CH02_FIG_05_Feature_Representations.png",
                4.2,
            ),
        },
        {
            "kicker": "Evidence selection",
            "title": "Review Strategy and Source Evaluation",
            "paragraphs": [
                (
                    "The supplied 2025 survey provides the organizing taxonomy for this "
                    "chapter and summarizes research published primarily between 2020 and "
                    "2022. Its conclusions are supplemented with primary papers that "
                    "introduced influential datasets or methods, including UCF-Crime, "
                    "RWF-2000, temporal-regularity autoencoders, future-frame "
                    "prediction, three-dimensional convolution, two-stream networks, robust "
                    "temporal feature-magnitude learning, and video transformers. The review "
                    "prioritizes peer-reviewed or author-released primary sources and records "
                    "the task, dataset, metric, and supervision setting associated with each "
                    "reported result.",
                    "[1], [2], [11]–[22]",
                ),
                (
                    "Performance values are not transferred blindly across papers. A result "
                    "is considered comparable only when the dataset split, annotation level, "
                    "input modality, metric definition, and evaluation protocol are "
                    "compatible. The review also separates claims supported by an available "
                    "primary source from claims supplied in a secondary summary, and any "
                    "accuracy or performance value reported in later chapters must be "
                    "grounded in a verifiable benchmark with a clearly identified source or "
                    "in a reproducible CrimeLens experiment.",
                    None,
                ),
            ],
            "figure": (
                "Figure 2.6",
                "Literature-selection and evidence-synthesis workflow",
                "CH02_FIG_06_Literature_Review_Workflow.png",
                7.0,
            ),
            "info": (
                "Evidence rule",
                "Every numeric result in the final book should be traceable to a primary "
                "paper or to a reproducible CrimeLens experiment. Dataset names and metric "
                "names alone are insufficient evidence.",
                GOLD,
                "SOURCE",
            ),
        },
        {
            "kicker": "Algorithm taxonomy",
            "title": "Overview of Algorithm Families",
            "section": "2.2",
            "paragraphs": [
                (
                    "Video anomaly and crime detection algorithms can be organized by the "
                    "kind of evidence they model. Discriminative classifiers learn a direct "
                    "boundary between known classes. Reconstruction and prediction models "
                    "learn regularity and use error as an anomaly signal. Weakly supervised "
                    "ranking methods infer abnormal temporal segments from coarse labels. "
                    "Object detectors localize semantically meaningful items. Recurrent, "
                    "three-dimensional convolutional, and attention-based models capture "
                    "temporal dependencies, while hybrid systems combine "
                    "several sources of evidence.",
                    "[1]",
                ),
                (
                    "No family dominates every operating condition. Real-time object "
                    "detection can provide interpretable bounding boxes but cannot alone "
                    "explain behavior. Reconstruction models can reveal unseen deviations "
                    "but may also reconstruct anomalies well. Three-dimensional networks "
                    "capture motion directly but require larger datasets and computational "
                    "budgets. Transformers model long-range context but are data hungry. "
                    "The algorithm choice must therefore be linked to the anomaly definition, "
                    "available supervision, latency budget, acceptable false-alarm rate, and "
                    "need for human interpretation.",
                    None,
                ),
            ],
            "table": {
                "number": "Table 2.5",
                "caption": "Algorithm families and principal trade-offs",
                "headers": ["Family", "Learning signal", "Primary output", "Typical trade-off"],
                "rows": [
                    ["2D CNN / classifier", "Class labels", "Frame or clip class", "Efficient but temporally limited"],
                    ["3D CNN / video model", "Clip labels", "Spatiotemporal class/features", "Higher memory and data demand"],
                    ["AE / prediction", "Normal reconstruction target", "Reconstruction/prediction error", "Error may not equal anomaly"],
                    ["MIL / weak supervision", "Video-level labels", "Temporal anomaly score", "Localization uncertainty"],
                    ["YOLO / detector", "Bounding boxes and classes", "Localized objects", "Objects do not establish intent"],
                    ["Transformer", "Supervised or self-supervised", "Contextual video representation", "Large-scale training requirements"],
                ],
                "widths": [3.2, 4.0, 4.2, 4.7],
                "accent": PURPLE,
                "font_size": 7.5,
            },
            "figure": (
                "Figure 2.7",
                "Taxonomy of algorithm families reviewed in Chapter Two",
                "CH02_FIG_07_Algorithm_Taxonomy.png",
                4.0,
            ),
        },
        {
            "kicker": "Spatial learning",
            "title": "Two-Dimensional Convolutional Neural Networks",
            "paragraphs": [
                (
                    "A two-dimensional convolutional neural network applies learned filters "
                    "across image height and width. Early layers capture edges and local "
                    "patterns, deeper layers combine them into objects and scene concepts, "
                    "and a classification head maps the representation to a probability "
                    "distribution. For surveillance tasks, a network may classify individual "
                    "frames, short frame samples, detected crops, or features pooled from "
                    "several frames. Transfer learning from a large image dataset is common "
                    "because labeled crime footage is limited.",
                    "[1]",
                ),
                (
                    "The main advantage of two-dimensional convolution is operational "
                    "efficiency. Mature backbones, pretrained weights, compact variants, and "
                    "hardware acceleration make frame-level inference practical. The main "
                    "limitation is that movement is represented indirectly. Independently "
                    "classified frames may miss the order, duration, and interaction that "
                    "distinguish a fight from ordinary contact or a fall from sitting. "
                    "Temporal sampling, score smoothing, recurrent layers, optical flow, or "
                    "late fusion can partially compensate for this limitation.",
                    None,
                ),
                (
                    "The supplied survey reports that architectural heterogeneity among "
                    "pretrained two-dimensional backbones had a statistically negligible "
                    "overall influence in the reviewed crowd-anomaly experiments. This does "
                    "not mean that all models are identical; it suggests that data quality, "
                    "task formulation, temporal modeling, and deployment constraints may "
                    "matter more than selecting the deepest available image backbone.",
                    "[1]",
                ),
            ],
            "figure": (
                "Figure 2.8",
                "Two-dimensional convolutional classification pipeline",
                "CH02_FIG_08_2D_CNN_Pipeline.png",
                6.4,
            ),
        },
        {
            "kicker": "Motion modeling",
            "title": "Optical Flow and Two-Stream Networks",
            "paragraphs": [
                (
                    "Optical flow estimates apparent pixel displacement between frames and "
                    "provides an explicit representation of motion. Two-stream networks "
                    "separate appearance and motion: a spatial stream processes RGB frames, "
                    "while a temporal stream processes stacks of optical-flow fields. Their "
                    "scores or features are then fused. Simonyan and Zisserman demonstrated "
                    "that the two streams carry complementary information for action "
                    "recognition, establishing an influential pattern for later video "
                    "architectures.",
                    "[15]",
                ),
                (
                    "The separation is useful in surveillance because the same appearance "
                    "can participate in different actions and similar motion can involve "
                    "different objects. Motion highlights direction, speed, and interaction, "
                    "while RGB preserves object identity and scene context. However, dense "
                    "optical-flow calculation can be expensive and sensitive to blur, "
                    "compression, illumination changes, and moving cameras. Flow may also "
                    "amplify irrelevant background motion. End-to-end learned motion modules "
                    "and three-dimensional convolution reduce dependence on an external flow "
                    "estimator, but they do not remove the need to reason about motion quality.",
                    None,
                ),
            ],
            "figure": (
                "Figure 2.9",
                "Two-stream appearance and optical-flow architecture",
                "CH02_FIG_09_Two_Stream_Optical_Flow.png",
                8.2,
            ),
            "info": (
                "Interpretation",
                "Optical flow describes how image content moves. It does not identify why "
                "the motion occurred, so appearance, object, and scene evidence remain "
                "necessary for a defensible alert.",
                CYAN_DARK,
                "MOTION",
            ),
        },
        {
            "kicker": "Joint spatial-temporal learning",
            "title": "Three-Dimensional Convolutional Neural Networks",
            "paragraphs": [
                (
                    "Three-dimensional convolution extends the filter across height, width, "
                    "and time. A kernel therefore observes a short video volume and can learn "
                    "motion-sensitive patterns directly rather than classifying each frame "
                    "independently. The C3D work showed that homogeneous 3 × 3 × 3 kernels "
                    "produce useful generic spatiotemporal features, while Inflated 3D "
                    "ConvNets expanded successful two-dimensional filters into three "
                    "dimensions and benefited strongly from pretraining on the large Kinetics "
                    "action dataset.",
                    "[13], [14]",
                ),
                (
                    "These models are attractive for crime and violence recognition because "
                    "many events are defined by motion and interaction. Their cost is higher: "
                    "clips consume more memory than frames, training requires more labeled "
                    "video, and inference latency grows with clip length, resolution, and "
                    "sampling rate. Short clips may miss long events, while long clips can "
                    "dilute brief evidence. Careful temporal windows, augmentation, transfer "
                    "learning, and balanced sampling are therefore essential.",
                    None,
                ),
            ],
            "table": {
                "number": "Table 2.6",
                "caption": "Comparison of two-dimensional and three-dimensional convolution",
                "headers": ["Criterion", "2D convolution", "3D convolution"],
                "rows": [
                    ["Input", "Individual frame or image crop", "Sequence / video volume"],
                    ["Learned evidence", "Primarily spatial appearance", "Joint spatial-temporal patterns"],
                    ["Data and memory", "Relatively lower", "Relatively higher"],
                    ["Temporal order", "Requires an external mechanism", "Represented inside the convolution"],
                    ["Typical use", "Objects, scenes, frame classes", "Actions, interactions, short events"],
                ],
                "widths": [4.2, 5.8, 5.8],
                "accent": PURPLE,
                "font_size": 8.0,
            },
            "figure": (
                "Figure 2.10",
                "Three-dimensional convolution over space and time",
                "CH02_FIG_10_3D_CNN_Video_Volume.png",
                4.0,
            ),
            "info": (
                "Evidence requirement",
                "Reported accuracy for a three-dimensional convolutional network depends "
                "heavily on the dataset, class balance, split, and metric definition. Any "
                "such figure adopted by CrimeLens must therefore rest on a verifiable "
                "benchmark with a clearly identified source or on a reproducible in-house "
                "experiment rather than on a secondary summary.",
                PURPLE,
            ),
        },
        {
            "kicker": "Sequential memory",
            "title": "Recurrent Networks and Temporal Sequence Modeling",
            "paragraphs": [
                (
                    "Recurrent neural networks update a hidden state as a sequence is "
                    "processed. Long Short-Term Memory networks introduce gated memory to "
                    "preserve useful information and reduce the vanishing-gradient problem. "
                    "In video analysis, a convolutional backbone commonly extracts one "
                    "feature vector per frame or segment, and an LSTM or bidirectional LSTM "
                    "then models temporal order. Convolutional LSTM variants retain spatial "
                    "structure inside the recurrent state and are therefore useful for "
                    "frame prediction and motion-aware representation.",
                    "[22]",
                ),
                (
                    "Recurrent modeling can integrate evidence across a longer interval than "
                    "a small three-dimensional convolutional window. It can represent event "
                    "build-up, persistence, and recovery, and it supports variable-length "
                    "sequences. Limitations include sequential computation, sensitivity to "
                    "sampling and sequence length, and the possibility that important brief "
                    "events are diluted by long normal regions. Bidirectional processing "
                    "uses future context and is valuable for offline analysis, but it is not "
                    "directly suitable for a strictly online alarm unless delay is accepted.",
                    None,
                ),
                (
                    "Temporal smoothing is a simpler related technique. Instead of learning "
                    "a complete recurrent model, it aggregates predictions across nearby "
                    "frames using moving averages, voting, hysteresis, or persistence rules. "
                    "Smoothing can reduce isolated false positives, but excessive smoothing "
                    "increases detection delay and may suppress short genuine events.",
                    None,
                ),
            ],
            "figure": (
                "Figure 2.11",
                "Recurrent temporal-modeling pipeline",
                "CH02_FIG_11_Recurrent_Temporal_Model.png",
                6.8,
            ),
        },
        {
            "kicker": "Learning normality",
            "title": "Autoencoder-Based Anomaly Detection",
            "paragraphs": [
                (
                    "An autoencoder contains an encoder that compresses the input and a "
                    "decoder that reconstructs it. When trained primarily on normal video, "
                    "the model is expected to reconstruct familiar patterns accurately and "
                    "produce a larger error for abnormal patterns. Hasan and colleagues "
                    "developed fully connected and fully convolutional autoencoder approaches "
                    "for learning temporal regularity with little supervision. The resulting "
                    "reconstruction error can be normalized into a regularity or anomaly "
                    "score.",
                    "[11]",
                ),
                (
                    "The method is conceptually attractive because abnormal examples are "
                    "rare and difficult to enumerate. Nevertheless, reconstruction error is "
                    "not a guaranteed semantic anomaly measure. A high-capacity decoder may "
                    "reconstruct abnormal inputs surprisingly well, while harmless changes "
                    "in illumination, perspective, or background can create large errors. "
                    "Pixel-level error also assigns importance to every image region, even "
                    "when only people or objects are operationally relevant.",
                    None,
                ),
                (
                    "Research has responded with memory modules, sparse latent spaces, "
                    "object-centric reconstruction, perceptual losses, and combinations of "
                    "appearance and motion. These modifications attempt to prevent the model "
                    "from simply copying the input and to make the anomaly score more aligned "
                    "with meaningful scene content.",
                    "[1]",
                ),
            ],
            "figure": (
                "Figure 2.12",
                "Reconstruction-based autoencoder anomaly scoring",
                "CH02_FIG_12_Autoencoder_Anomaly_Scoring.png",
                7.0,
            ),
        },
        {
            "kicker": "Predictive normality",
            "title": "Future-Frame Prediction",
            "paragraphs": [
                (
                    "Prediction-based methods learn to forecast a future frame or feature "
                    "representation from preceding observations. If the model has learned "
                    "normal appearance and motion, an unusual event should be more difficult "
                    "to predict. Liu and colleagues combined intensity, gradient, and optical-"
                    "flow constraints so that predicted normal frames remain visually and "
                    "temporally consistent. A prediction-quality measure is then converted "
                    "into an anomaly score.",
                    "[12]",
                ),
                (
                    "Prediction can be stricter than reconstruction because the model cannot "
                    "copy the target frame directly. It must infer what should occur next. "
                    "However, ordinary scenes are often multimodal: several future actions "
                    "may all be normal. A deterministic predictor may blur these possibilities "
                    "and create error even without an anomaly. Camera noise, sudden lighting "
                    "changes, and unpredictable but harmless motion also affect the score.",
                    None,
                ),
                (
                    "Convolutional recurrent and variational recurrent models extend the "
                    "approach by representing uncertainty and temporal state. The design "
                    "choice remains operational: the prediction horizon must be long enough "
                    "to capture meaningful motion but short enough to preserve image quality "
                    "and low-latency detection.",
                    "[22]",
                ),
            ],
            "figure": (
                "Figure 2.13",
                "Future-frame prediction and prediction-error scoring",
                "CH02_FIG_13_Future_Frame_Prediction.png",
                7.2,
            ),
        },
        {
            "kicker": "Adversarial learning",
            "title": "Generative Adversarial Approaches",
            "paragraphs": [
                (
                    "A Generative Adversarial Network trains a generator and discriminator "
                    "in competition. The generator attempts to produce outputs that resemble "
                    "the training distribution, while the discriminator attempts to "
                    "distinguish generated samples from real samples. In anomaly detection, "
                    "adversarial losses can sharpen reconstructed or predicted frames, learn "
                    "a manifold of normal behavior, or provide a discriminator-based anomaly "
                    "signal.",
                    "[17]",
                ),
                (
                    "Adversarial training addresses the blur often produced by pixel-wise "
                    "reconstruction losses, but it introduces instability, mode collapse, "
                    "sensitivity to hyperparameters, and difficulty interpreting the final "
                    "score. A visually realistic generated frame does not guarantee that the "
                    "model has understood the event, and the discriminator may focus on "
                    "texture rather than behavior. For surveillance, these risks are "
                    "important because score calibration and reproducibility matter as much "
                    "as visual quality.",
                    "[1]",
                ),
            ],
            "table": {
                "number": "Table 2.7",
                "caption": "Reconstruction, prediction, and adversarial approaches",
                "headers": ["Approach", "Training objective", "Anomaly signal", "Common weakness"],
                "rows": [
                    ["Reconstruction", "Rebuild normal input", "Reconstruction error", "Anomalies may also reconstruct well"],
                    ["Prediction", "Forecast expected future", "Prediction error", "Normal futures can be uncertain"],
                    ["Adversarial", "Match normal data distribution", "Generator/discriminator score", "Training instability and calibration"],
                    ["Hybrid", "Combine several objectives", "Weighted multi-score", "Complexity and difficult attribution"],
                ],
                "widths": [3.0, 4.5, 4.0, 4.7],
                "accent": RED,
                "font_size": 8.0,
            },
            "figure": (
                "Figure 2.14",
                "Adversarial learning for video anomaly detection",
                "CH02_FIG_14_GAN_Anomaly_Framework.png",
                4.5,
            ),
        },
        {
            "kicker": "Weak supervision",
            "title": "Multiple-Instance Learning and Anomaly Ranking",
            "paragraphs": [
                (
                    "Multiple-Instance Learning represents each long video as a bag of "
                    "temporal instances. A normal bag is expected to contain normal "
                    "instances, whereas an abnormal bag contains at least one abnormal "
                    "instance, although its position is unknown during training. Sultani, "
                    "Chen, and Shah used this formulation to learn an anomaly-ranking model "
                    "from video-level labels. The loss encourages the highest-scoring "
                    "instance in an abnormal bag to exceed the highest-scoring instance in a "
                    "normal bag by a margin.",
                    "[2]",
                ),
                (
                    "Sparsity reflects the assumption that anomalies occupy a limited part "
                    "of a long video, while temporal smoothness discourages implausibly large "
                    "score changes between adjacent segments. Together, these constraints "
                    "improve temporal localization without frame-level training labels. The "
                    "approach significantly reduced annotation cost and established "
                    "UCF-Crime as a central benchmark for real-world weakly supervised video "
                    "anomaly detection.",
                    "[2]",
                ),
                (
                    "The formulation also has limitations. A positive bag can contain many "
                    "normal instances that dominate training; the highest score may lock onto "
                    "an easy but irrelevant visual cue; and coarse segments limit boundary "
                    "precision. Dataset bias can cause the model to associate camera style, "
                    "scene, or compression artifacts with the positive label rather than the "
                    "event itself.",
                    None,
                ),
            ],
            "figure": (
                "Figure 2.15",
                "Multiple-instance learning with video bags and temporal instances",
                "CH02_FIG_15_MIL_Bags_Instances.png",
                7.0,
            ),
        },
        {
            "kicker": "Improved weak supervision",
            "title": "Robust Temporal Feature-Magnitude Learning",
            "paragraphs": [
                (
                    "Robust Temporal Feature Magnitude learning addresses the imbalance "
                    "inside abnormal bags, where rare positive instances are surrounded by "
                    "many normal segments. Instead of relying only on anomaly scores, the "
                    "method learns feature magnitudes that separate abnormal and normal "
                    "instances. Dilated temporal convolutions and self-attention capture "
                    "short- and long-range dependencies, while the learning objective "
                    "emphasizes hard examples.",
                    "[19]",
                ),
                (
                    "The method illustrates a broader trend in weakly supervised research: "
                    "performance depends not only on the video backbone but also on how "
                    "instances are selected, ranked, aggregated, and regularized. Segment "
                    "features must preserve temporal evidence; the loss must resist dominant "
                    "negative instances; and the evaluation must distinguish video-level "
                    "classification from frame-level localization. Reported improvements "
                    "across UCF-Crime, ShanghaiTech, and UCSD-Peds show the value "
                    "of designing the learning objective around weak-label noise.",
                    "[19]",
                ),
            ],
            "figure": (
                "Figure 2.16",
                "Robust temporal feature-magnitude learning",
                "CH02_FIG_16_RTFM_Feature_Magnitude.png",
                8.3,
            ),
            "info": (
                "Research lesson",
                "In weak supervision, the loss function and temporal aggregation strategy "
                "can be as important as the visual backbone because they determine which "
                "unlabeled segments receive learning pressure.",
                PURPLE,
                "MIL",
            ),
        },
        {
            "kicker": "Localized semantic evidence",
            "title": "Object Detection and YOLO-Based Methods",
            "paragraphs": [
                (
                    "Object detection predicts both class probabilities and bounding boxes. "
                    "The original You Only Look Once architecture framed detection as a "
                    "single regression problem from the full image to spatially separated "
                    "boxes and classes, enabling real-time processing. For surveillance, a "
                    "detector can localize weapons, fire, people, vehicles, or protected "
                    "objects and produce visual evidence that is easier for an operator to "
                    "inspect than a single unexplained anomaly score.",
                    "[16]",
                ),
                (
                    "Detection is nevertheless only one level of understanding. A knife in "
                    "a kitchen, a firearm carried by an authorized officer, or a fast-moving "
                    "vehicle on a road may be visually detectable without being anomalous in "
                    "context. Small objects, occlusion, low resolution, motion blur, unusual "
                    "viewpoints, and class imbalance further affect performance. Thresholds "
                    "and non-maximum suppression control the final boxes but also trade recall "
                    "against false alerts.",
                    None,
                ),
                (
                    "Research systems often combine detection with tracking, pose, action "
                    "recognition, zones, trajectories, or temporal persistence. This "
                    "object-centric structure supports interpretable rules and relational "
                    "reasoning, but errors propagate: a missed object cannot be recovered by "
                    "later logic, and an identity switch can corrupt interaction features.",
                    "[1]",
                ),
            ],
            "figure": (
                "Figure 2.17",
                "Object-detection and scene-evidence pipeline",
                "CH02_FIG_17_YOLO_Object_Evidence.png",
                7.0,
            ),
        },
        {
            "kicker": "Long-range context",
            "title": "Attention and Video Transformers",
            "paragraphs": [
                (
                    "Self-attention computes relationships among elements of a sequence and "
                    "allows a representation to emphasize relevant spatial or temporal "
                    "evidence. TimeSformer applies attention to video patches and found that "
                    "separating spatial and temporal attention inside each block provided an "
                    "effective design. ViViT similarly represents video as spatiotemporal "
                    "tokens and studies factorized variants that reduce the cost of very long "
                    "token sequences.",
                    "[20], [21]",
                ),
                (
                    "Transformers are attractive for surveillance because a significant "
                    "event may depend on relationships separated by many frames or distant "
                    "image regions. Attention can connect a person to an object, relate an "
                    "early approach to a later interaction, and integrate longer clips than "
                    "a small convolutional kernel. However, global attention is expensive, "
                    "pretraining data are important, and the learned attention weights are "
                    "not automatically a faithful causal explanation.",
                    None,
                ),
                (
                    "Efficient variants use divided attention, local windows, sparse tokens, "
                    "feature pyramids, or convolutional front ends. For anomaly detection, "
                    "the challenge is not simply adopting a transformer but defining the "
                    "training signal, token granularity, temporal resolution, and anomaly "
                    "score in a way that survives long normal intervals and rare short events.",
                    "[1]",
                ),
            ],
            "figure": (
                "Figure 2.18",
                "Space-time attention for video understanding",
                "CH02_FIG_18_Video_Transformer_Attention.png",
                7.0,
            ),
        },
        {
            "kicker": "Integrated inference",
            "title": "Hybrid Models and Evidence Fusion",
            "paragraphs": [
                (
                    "Hybrid methods combine complementary inductive biases. A detector can "
                    "localize people and objects; a tracker can maintain identity; a pose or "
                    "action model can describe motion; a recurrent or attention module can "
                    "aggregate time; and a rule or learned fusion layer can combine the "
                    "scores. Reconstruction error can be fused with discriminative "
                    "probability from a detector or action model. The "
                    "motivation is that no single representation captures object identity, "
                    "movement, interaction, context, and event duration equally well.",
                    "[1]",
                ),
                (
                    "Fusion has several forms. Feature-level fusion learns interactions but "
                    "requires compatible scales and synchronized inputs. Score-level fusion "
                    "is simpler and allows independent models, but calibration is essential "
                    "because a value of 0.8 may have different meanings across classifiers. "
                    "Decision-level fusion uses explicit rules, voting, persistence, or "
                    "priority logic and is easier to explain, although fixed rules can fail "
                    "under new conditions.",
                    None,
                ),
                (
                    "Hybrid complexity must be justified by measured benefit. Each additional "
                    "model introduces latency, resource use, failure states, monitoring, and "
                    "calibration work. Ablation studies are therefore necessary to show which "
                    "components improve recall, reduce false alarms, or increase robustness "
                    "rather than merely increasing architectural size.",
                    None,
                ),
            ],
            "figure": (
                "Figure 2.19",
                "Hybrid evidence-fusion strategies",
                "CH02_FIG_19_Hybrid_Evidence_Fusion.png",
                7.2,
            ),
        },
        {
            "kicker": "Comparison protocol",
            "title": "Principles for Comparing Algorithms",
            "section": "2.3",
            "paragraphs": [
                (
                    "A fair algorithm comparison begins by matching the task. Frame-level "
                    "anomaly localization, video-level classification, temporal event "
                    "localization, object detection, and multi-class crime recognition have "
                    "different outputs. A method evaluated on short balanced clips cannot be "
                    "directly ranked against one evaluated on long untrimmed video. Likewise, "
                    "accuracy on a binary violence dataset is not equivalent to frame-level "
                    "Area Under the Receiver Operating Characteristic Curve on an anomaly "
                    "dataset.",
                    None,
                ),
                (
                    "The comparison must record the dataset version and split, supervision "
                    "level, feature extractor, input resolution, temporal window, modality, "
                    "pretraining, and whether external data were used. Runtime evidence also "
                    "matters: frames per second depends on hardware, batch size, decoding, "
                    "preprocessing, and whether optical flow was precomputed. Model parameter "
                    "count alone does not describe end-to-end latency.",
                    None,
                ),
                (
                    "Statistical uncertainty should be acknowledged. Random initialization, "
                    "class imbalance, threshold selection, and limited test events can change "
                    "results. Multiple runs, confidence intervals, per-class analysis, and "
                    "cross-dataset tests provide stronger evidence than a single aggregate "
                    "number. The supplied survey’s nonparametric analysis is valuable because "
                    "it evaluates whether observed differences among pretrained backbones "
                    "are systematic rather than merely numerical.",
                    "[1]",
                ),
            ],
            "info": (
                "Comparison checklist",
                "Task · dataset and split · annotation level · metric · modality · "
                "pretraining · threshold · hardware · latency boundary · number of runs.",
                CYAN_DARK,
                "FAIR",
            ),
        },
        {
            "kicker": "Measurement",
            "title": "Evaluation Metrics and Operating Thresholds",
            "paragraphs": [
                (
                    "Precision measures how many reported positives are correct, while recall "
                    "measures how many true positives are detected. Their harmonic mean, the "
                    "F1 score, balances both but still depends on a chosen threshold. The "
                    "Receiver Operating Characteristic curve plots true-positive rate against "
                    "false-positive rate across thresholds, and its Area Under the Curve "
                    "summarizes ranking performance. Precision–Recall curves are often more "
                    "informative when anomalies are rare because they emphasize positive-class "
                    "performance.",
                    "[1]",
                ),
                (
                    "Frame-level and pixel-level AUC describe different localization scales. "
                    "Equal Error Rate identifies the operating point where false acceptance "
                    "and false rejection are equal. Detection delay measures how long the "
                    "system takes to raise an alert after event onset. Frames per second and "
                    "end-to-end latency describe throughput and responsiveness. A practical "
                    "surveillance evaluation should report false alarms over time or per "
                    "camera because even a small frame-level false-positive rate can produce "
                    "an unmanageable number of alerts in continuous operation.",
                    None,
                ),
            ],
            "table": {
                "number": "Table 2.9",
                "caption": "Evaluation metrics and interpretation",
                "headers": ["Metric", "Measures", "Useful when", "Important caution"],
                "rows": [
                    ["Precision", "Reliability of positive alerts", "False alerts are costly", "Can rise by missing difficult events"],
                    ["Recall", "Coverage of true events", "Missed events are costly", "Can rise by alerting too often"],
                    ["F1", "Balance of precision and recall", "One threshold must be summarized", "Hides operating alternatives"],
                    ["ROC-AUC", "Ranking over all thresholds", "Comparing score-based models", "May appear optimistic under imbalance"],
                    ["PR-AUC", "Positive-class ranking", "Anomalies are rare", "Depends strongly on prevalence"],
                    ["Delay / FPS", "Operational responsiveness", "Online deployment", "Hardware and pipeline boundaries matter"],
                ],
                "widths": [2.5, 4.0, 4.6, 5.1],
                "accent": GREEN,
                "font_size": 7.4,
            },
            "figure": (
                "Figure 2.20",
                "Evaluation metrics and threshold behavior",
                "CH02_FIG_20_Metrics_Thresholds.png",
                3.9,
            ),
        },
        {
            "kicker": "Benchmark landscape",
            "title": "Datasets for Anomaly and Violence Detection",
            "paragraphs": [
                (
                    "Benchmark datasets differ in realism, duration, camera count, scene "
                    "diversity, anomaly categories, annotation granularity, and modality. "
                    "UCSD Pedestrian and CUHK Avenue are compact scene-specific benchmarks "
                    "that support frame- or pixel-level localization. ShanghaiTech contains "
                    "multiple campus scenes and greater environmental variation. UCF-Crime "
                    "contains long real-world surveillance videos and video-level labels for "
                    "13 anomaly categories. RWF-2000 focuses on binary violence clips drawn "
                    "from real surveillance footage.",
                    "[1], [2], [18]",
                ),
                (
                    "A dataset is not simply a collection size. Staged versus real events, "
                    "trimmed versus untrimmed clips, camera viewpoint, compression, class "
                    "balance, temporal boundaries, and train-test leakage all affect the "
                    "difficulty. Models can exploit background and source artifacts when "
                    "classes are collected differently. Consequently, dataset documentation "
                    "and split discipline are essential to interpreting results.",
                    None,
                ),
            ],
            "table": {
                "number": "Table 2.8",
                "caption": "Representative benchmark datasets",
                "headers": ["Dataset", "Scale / setting", "Labels", "Primary task"],
                "rows": [
                    ["UCSD Ped1/Ped2", "Fixed pedestrian scenes", "Frame/pixel anomaly", "Normality-based detection"],
                    ["CUHK Avenue", "16 training and 21 test clips", "Frame/pixel anomaly", "Scene anomaly detection"],
                    ["ShanghaiTech", "13 campus scenes", "Frame anomaly", "Multi-scene generalization"],
                    ["UCF-Crime", "1,900 videos; about 128 hours", "Video labels; test temporal annotations", "Weakly supervised anomaly detection"],
                    ["RWF-2000", "2,000 real surveillance clips", "Violent / non-violent", "Binary violence recognition"],
                ],
                "widths": [2.8, 4.7, 4.2, 4.5],
                "accent": CYAN_DARK,
                "font_size": 7.4,
            },
            "figure": (
                "Figure 2.21",
                "Dataset landscape by realism, duration, and annotation level",
                "CH02_FIG_21_Dataset_Landscape.png",
                3.8,
            ),
        },
        {
            "kicker": "Large-scale anomaly benchmark",
            "title": "UCF-Crime Dataset",
            "paragraphs": [
                (
                    "UCF-Crime was introduced to move anomaly detection toward realistic, "
                    "long, untrimmed surveillance footage. The dataset contains 1,900 videos "
                    "with approximately 128 hours of content. Thirteen anomaly categories are "
                    "represented: abuse, arrest, arson, assault, burglary, explosion, "
                    "fighting, road accident, robbery, shooting, shoplifting, stealing, and "
                    "vandalism, together with normal activities. The dataset supports general "
                    "anomaly detection and category recognition.",
                    "[2]",
                ),
                (
                    "Its weak training labels reflect a practical annotation setting: the "
                    "video is labeled normal or anomalous, but exact event boundaries are not "
                    "provided to the learner. Long normal regions inside positive videos make "
                    "the problem difficult and motivate multiple-instance learning. The "
                    "dataset also exposes models to diverse scenes, viewpoints, video quality, "
                    "and event duration. These properties improve realism but introduce "
                    "background bias and temporal-localization uncertainty.",
                    None,
                ),
                (
                    "UCF-Crime is frequently used for frame-level AUC evaluation after "
                    "producing a temporal anomaly score. A complete report should state the "
                    "feature extractor, segment count, training split, smoothing, and whether "
                    "the evaluation follows the original protocol. Multi-class recognition "
                    "results should not be confused with general anomaly-detection results.",
                    None,
                ),
            ],
            "figure": (
                "Figure 2.22",
                "UCF-Crime categories and weak labels",
                "CH02_FIG_22_UCF_Crime_Categories.png",
                7.2,
            ),
        },
        {
            "kicker": "Compact benchmarks",
            "title": "UCSD, Avenue, and ShanghaiTech",
            "paragraphs": [
                (
                    "UCSD Ped1 and Ped2 contain fixed-camera pedestrian scenes in which "
                    "bicycles, vehicles, skaters, or unusual movement can be anomalous. Their "
                    "controlled setting and pixel-level annotations made them important for "
                    "evaluating reconstruction and localization methods. CUHK Avenue adds "
                    "varied actions, camera shake, and events such as running, throwing, and "
                    "unusual object movement. ShanghaiTech expands the problem to 13 scenes "
                    "with different viewpoints, lighting, density, and anomaly types.",
                    "[1]",
                ),
                (
                    "These datasets are valuable because they support detailed temporal or "
                    "spatial evaluation, but their scale and scene distribution differ from "
                    "long operational crime footage. A model may learn the normality of one "
                    "fixed scene and perform well without developing general semantic crime "
                    "understanding. ShanghaiTech reduces this weakness through multiple "
                    "scenes, although domain shift remains. Reported scores should therefore "
                    "be interpreted as benchmark-specific evidence rather than universal "
                    "surveillance accuracy.",
                    None,
                ),
                (
                    "The three benchmarks are particularly useful for studying anomaly-score "
                    "quality, object-centric reconstruction, future-frame prediction, memory "
                    "models, and cross-scene robustness. They complement rather than replace "
                    "large-scale real-world datasets.",
                    None,
                ),
            ],
            "figure": (
                "Figure 2.23",
                "Representative compact anomaly-detection benchmarks",
                "CH02_FIG_23_Compact_Benchmark_Scenes.jpg",
                7.5,
            ),
        },
        {
            "kicker": "Violence benchmark",
            "title": "RWF-2000 Violence Benchmark",
            "paragraphs": [
                (
                    "RWF-2000 contains 2,000 real-world surveillance clips divided between "
                    "violent and non-violent behavior. Its authors proposed a Flow Gated "
                    "Network combining three-dimensional convolution and optical flow and "
                    "reported 87.25% accuracy on the dataset’s test set. The benchmark is "
                    "useful for binary violence recognition because its footage originates "
                    "from surveillance scenes rather than only movies or sports.",
                    "[18]",
                ),
                (
                    "RWF-2000 supports a balanced clip-level violent/non-violent "
                    "classification protocol. Because its clips are short and trimmed, its "
                    "reported values must not be compared directly with weakly supervised "
                    "anomaly-localization results on long untrimmed videos such as UCF-Crime, "
                    "where the task, label granularity, and metric definitions differ.",
                    None,
                ),
            ],
            "figure": (
                "Figure 2.24",
                "RWF-2000 violence dataset characteristics",
                "CH02_FIG_24_Violence_Datasets.png",
                7.0,
            ),
        },
        {
            "kicker": "External validity",
            "title": "Cross-Dataset Generalization and Domain Shift",
            "paragraphs": [
                (
                    "A model can achieve a strong result on a benchmark and still fail on a "
                    "new camera. Domain shift arises from viewpoint, focal length, resolution, "
                    "frame rate, compression, illumination, weather, scene layout, crowd "
                    "density, clothing, cultural behavior, and event prevalence. Collection "
                    "bias can also link labels to source-specific artifacts. When training and "
                    "test videos share backgrounds or acquisition pipelines, the score may "
                    "overestimate semantic generalization.",
                    "[1]",
                ),
                (
                    "Cross-dataset testing provides stronger evidence by training on one "
                    "distribution and evaluating on another. Domain adaptation, feature "
                    "normalization, augmentation, self-supervised pretraining, and continual "
                    "calibration can reduce the gap, but they introduce assumptions and may "
                    "require target-domain data. Open-set evaluation is also important because "
                    "operational anomalies are not limited to categories present during "
                    "training.",
                    None,
                ),
                (
                    "Generalization should be measured at several levels: new videos from "
                    "known cameras, new cameras in known environments, new environments, and "
                    "new event types. Reporting only a random clip split cannot establish all "
                    "four. False-alarm analysis by camera and environmental condition helps "
                    "identify whether failure is caused by the event model or by the imaging "
                    "domain.",
                    None,
                ),
            ],
            "figure": (
                "Figure 2.25",
                "Sources of cross-dataset domain shift",
                "CH02_FIG_25_Domain_Shift.png",
                7.0,
            ),
        },
        {
            "kicker": "Reported evidence",
            "title": "Performance Comparison Across Representative Studies",
            "paragraphs": [
                (
                    "The selected results below illustrate why metric and dataset context "
                    "must remain attached to every value. RWF-2000 reports test accuracy for "
                    "balanced violence clips, while UCF-Crime studies weakly supervised "
                    "anomaly localization in long videos. Action-recognition results on "
                    "UCF-101 or HMDB-51 measure known action classes rather than surveillance "
                    "anomaly detection. A reported accuracy is therefore meaningful only "
                    "alongside its dataset, split, class balance, and metric definition, and "
                    "any value without a verifiable source is excluded from this comparison.",
                    "[2], [13], [14], [18]",
                ),
            ],
            "table": {
                "number": "Table 2.10",
                "caption": "Reported results selected from primary sources",
                "headers": ["Study / method", "Dataset / task", "Reported result", "Interpretation boundary"],
                "rows": [
                    ["Sultani et al. MIL", "UCF-Crime anomaly detection", "Significant improvement over compared baselines", "Use the paper’s frame-level protocol; not clip accuracy"],
                    ["C3D", "UCF-101 action recognition", "52.8% with compact learned features", "Action recognition; not anomaly detection"],
                    ["I3D", "UCF-101 / HMDB-51", "98.0% / 80.9%", "Kinetics-pretrained action recognition"],
                    ["Flow Gated Network", "RWF-2000 violence classification", "87.25% accuracy", "Balanced binary surveillance clips"],
                ],
                "widths": [3.8, 4.1, 3.7, 4.7],
                "accent": GOLD,
                "font_size": 7.2,
            },
            "figure": (
                "Figure 2.26",
                "Chronological evolution of representative literature",
                "CH02_FIG_26_Literature_Timeline.png",
                4.2,
            ),
            "info": (
                "Do not rank unlike metrics",
                "Accuracy, frame-level ROC-AUC, PR-AUC, and action-recognition accuracy "
                "answer different questions. A larger number is not automatically a better "
                "surveillance system.",
                RED,
                "METRIC",
            ),
        },
        {
            "kicker": "Synthesis and gaps",
            "title": "Literature Survey and Research Opportunities",
            "section": "2.4",
            "paragraphs": [
                (
                    "The literature shows a transition from handcrafted motion descriptors "
                    "to learned spatial-temporal representations, from small fixed-scene "
                    "benchmarks to long real-world videos, and from frame reconstruction to "
                    "weak supervision and attention. Despite this "
                    "progress, the core problem remains open. Anomalies are contextual, "
                    "labeled events are scarce, performance is dataset dependent, and a model "
                    "score does not by itself establish operational meaning.",
                    "[1]",
                ),
            ],
            "table": {
                "number": "Table 2.11",
                "caption": "Literature-survey synthesis",
                "headers": ["Research direction", "Strength", "Persistent gap"],
                "rows": [
                    ["Reconstruction / prediction", "Learns from normal video", "Novel normal events and over-general reconstruction"],
                    ["Weak supervision", "Reduces temporal annotation cost", "Noisy instance discovery and coarse boundaries"],
                    ["3D CNN / recurrent video models", "Captures temporal dynamics", "Data, compute, and window-length constraints"],
                    ["Object-centric detection", "Localized and interpretable evidence", "Intent and context remain unresolved"],
                    ["Transformers / attention", "Long-range relationships", "Scale, efficiency, and explanation limits"],
                ],
                "widths": [4.0, 5.2, 7.0],
                "accent": CYAN_DARK,
                "font_size": 7.5,
            },
            "figure": (
                "Figure 2.27",
                "Research-gap and opportunity map",
                "CH02_FIG_27_Research_Gap_Map.png",
                5.0,
            ),
            "info": (
                "Chapter conclusion",
                "The strongest research direction is not unrestricted automation. It is "
                "context-aware, calibrated, explainable, privacy-conscious detection whose "
                "output can be validated within a governed operational workflow.",
                GREEN,
                "SYNTHESIS",
            ),
        },
    ]


def render_chapter_two_page(document: Document, spec: dict, *, is_last: bool) -> None:
    add_page_heading(
        document,
        spec["kicker"],
        spec["title"],
        spec.get("section"),
    )
    for text, citation in spec.get("paragraphs", []):
        add_body(document, text, citation=citation)

    table = spec.get("table")
    if table:
        add_table_caption(document, table["number"], table["caption"])
        add_professional_table(
            document,
            table["headers"],
            table["rows"],
            table.get("widths"),
            accent=table.get("accent", CYAN_DARK),
            font_size=table.get("font_size", 8.0),
        )

    table_2 = spec.get("table_2")
    if table_2:
        add_table_caption(document, table_2["number"], table_2["caption"])
        add_professional_table(
            document,
            table_2["headers"],
            table_2["rows"],
            table_2.get("widths"),
            accent=table_2.get("accent", CYAN_DARK),
            font_size=table_2.get("font_size", 8.0),
        )

    for text, citation in spec.get("post_table_paragraphs", []):
        add_body(document, text, citation=citation)

    code_snippet = spec.get("code")
    if code_snippet:
        add_code_block(document, code_snippet[0], code_snippet[1])

    code_snippet_2 = spec.get("code_2")
    if code_snippet_2:
        add_code_block(document, code_snippet_2[0], code_snippet_2[1])

    figure = spec.get("figure")
    if figure:
        add_figure_placeholder(
            document,
            figure[0],
            figure[1],
            figure[2],
            figure[3],
        )

    info = spec.get("info")
    if info:
        if len(info) > 3:
            add_info_box(document, info[0], info[1], info[2], info[3])
        else:
            add_info_box(document, info[0], info[1], info[2])

    bullets = spec.get("bullets")
    if bullets:
        add_bullets(document, bullets)

    add_chapter_page_break(document, is_last=is_last)


def add_chapter_two_references(document: Document) -> None:
    add_page_break(document)
    heading = document.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run("Chapter Two References")

    for index, (number, reference) in enumerate(CHAPTER_TWO_REFERENCES):
        if index == 7:
            add_page_break(document)
            heading = document.add_paragraph(style="Heading 1")
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.add_run("Chapter Two References — Continued")

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.first_line_indent = Cm(-0.8)
        paragraph.paragraph_format.space_after = Pt(8)
        number_run = paragraph.add_run(number + " ")
        set_run_font(number_run, "Cambria", 9.5, NAVY, True)
        reference_run = paragraph.add_run(reference)
        set_run_font(reference_run, "Cambria", 9.5, TEXT)


def add_chapter_three_opener(document: Document) -> None:
    add_image_chapter_cover(
        document,
        "chapter3.png",
        "CHAPTER THREE",
        "SYSTEM ANALYSIS",
        "Actors  ·  Requirements  ·  DFDs  ·  Use Cases  ·  Sequences  ·  "
        "Activities  ·  State Machines  ·  Conceptual Classes",
        "This chapter defines what CrimeLens must do, who is authorized to do it, "
        "how information and control move, and which states and domain concepts govern "
        "the complete detection-to-response workflow.",
        title_size=40,
    )


def add_chapter_three_references(document: Document) -> None:
    add_page_break(document)
    heading = document.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run("Chapter Three References")

    for number, reference in CHAPTER_THREE_REFERENCES:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.first_line_indent = Cm(-0.8)
        paragraph.paragraph_format.space_after = Pt(10)
        number_run = paragraph.add_run(number + " ")
        set_run_font(number_run, "Cambria", 9.5, NAVY, True)
        reference_run = paragraph.add_run(reference)
        set_run_font(reference_run, "Cambria", 9.5, TEXT)


def add_chapter_four_opener(document: Document) -> None:
    add_image_chapter_cover(
        document,
        "chapter4.png",
        "CHAPTER FOUR",
        "SYSTEM DESIGN",
        "Architecture  ·  Modules  ·  ERD  ·  Security  ·  Realtime  ·  "
        "Interfaces  ·  Queues  ·  Evidence Storage",
        "This chapter converts the system-analysis models into a practical design "
        "blueprint covering architecture, bounded contexts, data structures, security "
        "boundaries, user interfaces, integrations, and operational support.",
        title_size=42,
    )


def add_chapter_four_references(document: Document) -> None:
    add_page_break(document)
    heading = document.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run("Chapter Four References")

    for number, reference in CHAPTER_FOUR_REFERENCES:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.first_line_indent = Cm(-0.8)
        paragraph.paragraph_format.space_after = Pt(10)
        number_run = paragraph.add_run(number + " ")
        set_run_font(number_run, "Cambria", 9.5, NAVY, True)
        reference_run = paragraph.add_run(reference)
        set_run_font(reference_run, "Cambria", 9.5, TEXT)


def add_chapter_five_opener(document: Document) -> None:
    add_image_chapter_cover(
        document,
        "chapter5.png",
        "CHAPTER FIVE",
        "SYSTEM IMPLEMENTATION",
        "Laravel  ·  Inertia/React  ·  Flutter  ·  AI Model  ·  Gateway  ·  "
        "Realtime  ·  Queues  ·  Security  ·  Testing",
        "This chapter documents how the designed CrimeLens system was implemented "
        "as maintainable software: modular backend, interactive web consoles, field "
        "mobile application, AI runtime, streaming gateway, background workers, "
        "security controls, and repeatable development evidence.",
        title_size=34,
    )


def add_chapter_five_references(document: Document) -> None:
    add_page_break(document)
    heading = document.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run("Chapter Five References")

    for number, reference in CHAPTER_FIVE_REFERENCES:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.first_line_indent = Cm(-0.8)
        paragraph.paragraph_format.space_after = Pt(10)
        number_run = paragraph.add_run(number + " ")
        set_run_font(number_run, "Cambria", 9.5, NAVY, True)
        reference_run = paragraph.add_run(reference)
        set_run_font(reference_run, "Cambria", 9.5, TEXT)


def add_chapter_six_opener(document: Document) -> None:
    add_image_chapter_cover(
        document,
        "chapter6.png",
        "CHAPTER SIX",
        "SYSTEM TESTING AND EVALUATION",
        "Pest  ·  Validation  ·  Security  ·  Realtime  ·  Browser  ·  Flutter  ·  "
        "AI Metrics  ·  Streaming  ·  Performance",
        "This chapter evaluates CrimeLens through executed backend tests, "
        "traceability, validation evidence, security checks, realtime behavior, "
        "AI metrics, streaming readiness, performance methodology, and final "
        "evidence artifacts for project review.",
        title_size=28,
    )


def add_chapter_six_references(document: Document) -> None:
    add_page_break(document)
    heading = document.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run("Chapter Six References")

    for number, reference in CHAPTER_SIX_REFERENCES:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.first_line_indent = Cm(-0.8)
        paragraph.paragraph_format.space_after = Pt(10)
        number_run = paragraph.add_run(number + " ")
        set_run_font(number_run, "Cambria", 9.5, NAVY, True)
        reference_run = paragraph.add_run(reference)
        set_run_font(reference_run, "Cambria", 9.5, TEXT)


def add_chapter_seven_opener(document: Document) -> None:
    add_image_chapter_cover(
        document,
        "chapter7.png",
        "CHAPTER SEVEN",
        "CONCLUSION AND FUTURE WORK",
        "Achievements  ·  Contributions  ·  Lessons Learned  ·  Limitations  ·  "
        "Roadmap  ·  Ethics  ·  Final Defense",
        "This chapter summarizes what CrimeLens achieved, explains the technical "
        "and academic value of the work, identifies responsible limitations, and "
        "defines a practical roadmap for turning the graduation prototype into a "
        "more scalable, secure, measurable, and field-ready system.",
        title_size=30,
    )


def add_chapter_seven_references(document: Document) -> None:
    add_page_break(document)
    heading = document.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run("Chapter Seven References")

    for number, reference in CHAPTER_SEVEN_REFERENCES:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.first_line_indent = Cm(-0.8)
        paragraph.paragraph_format.space_after = Pt(10)
        number_run = paragraph.add_run(number + " ")
        set_run_font(number_run, "Cambria", 9.5, NAVY, True)
        reference_run = paragraph.add_run(reference)
        set_run_font(reference_run, "Cambria", 9.5, TEXT)


def build() -> None:
    HERE.mkdir(parents=True, exist_ok=True)
    document = Document()
    setup_styles(document)
    set_document_background(document)
    enable_update_fields(document)

    properties = document.core_properties
    properties.title = PROJECT_TITLE
    properties.subject = (
        "Graduation Project Documentation — Front Matter and Chapters One to Seven"
    )
    properties.author = "CrimeLens Graduation Project Team"
    properties.keywords = (
        "CrimeLens, AI, Computer-Aided Dispatch, Smart Surveillance, Graduation Project"
    )
    properties.comments = (
        "Editable Word source generated for Beni-Suef University, academic year 2025/2026."
    )

    add_cover(document)

    front_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_a4(front_section)
    unlink_section(front_section)
    set_page_number_format(front_section, "decimal", 1)
    # No running header (it collided with the page-border frame); footer carries the info.
    add_footer(front_section, mono=True)

    add_team_pages(document)
    add_acknowledgement_page(document)
    add_abstract_page(document)
    add_toc_page(document)
    add_list_page(document, "List of Figures", "Figure Caption", ALL_FIGURES)
    add_list_page(document, "List of Tables", "Table Caption", ALL_TABLES)

    chapter_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_a4(chapter_section)
    unlink_section(chapter_section)
    set_page_number_format(chapter_section, "decimal")
    # No running header for Chapter One — the scope title lives only in the footer.
    add_scope_footer(chapter_section, "Introduction Scope")

    reset_abbreviation_tracking()
    add_chapter_opener(document)
    chapter_page_02(document)
    chapter_page_03(document)
    chapter_page_04(document)
    chapter_page_05(document)
    chapter_page_06(document)
    chapter_page_07(document)
    chapter_page_08(document)
    chapter_page_09(document)
    chapter_page_10(document)
    chapter_page_11(document)
    chapter_page_12(document)
    chapter_page_13(document)
    chapter_page_14(document)
    chapter_page_15(document)
    chapter_page_16(document)
    chapter_page_17(document)
    chapter_page_18(document)
    chapter_page_19(document)
    chapter_page_20(document)
    chapter_page_21(document)
    chapter_page_22(document)
    chapter_page_23(document)
    chapter_page_24(document)
    chapter_page_25(document)
    add_references(document)

    chapter_two_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_a4(chapter_two_section)
    unlink_section(chapter_two_section)
    set_page_number_format(chapter_two_section, "decimal")
    # No running header for Chapter Two — the scope title lives only in the footer.
    add_scope_footer(chapter_two_section, "Literature Review Scope")

    reset_abbreviation_tracking()
    add_chapter_two_opener(document)
    chapter_two_specs = chapter_two_page_specs()
    for index, spec in enumerate(chapter_two_specs):
        render_chapter_two_page(
            document,
            spec,
            is_last=index == len(chapter_two_specs) - 1,
        )
    add_chapter_two_references(document)

    chapter_three_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_a4(chapter_three_section)
    unlink_section(chapter_three_section)
    set_page_number_format(chapter_three_section, "decimal")
    # No running header — the scope title lives only in the footer.
    add_scope_footer(chapter_three_section, "System Analysis Scope")

    reset_abbreviation_tracking()
    add_chapter_three_opener(document)
    chapter_three_specs = chapter_three_page_specs()
    for index, spec in enumerate(chapter_three_specs):
        render_chapter_two_page(
            document,
            spec,
            is_last=index == len(chapter_three_specs) - 1,
        )
    add_chapter_three_references(document)

    chapter_four_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_a4(chapter_four_section)
    unlink_section(chapter_four_section)
    set_page_number_format(chapter_four_section, "decimal")
    # No running header — the scope title lives only in the footer.
    add_scope_footer(chapter_four_section, "System Design Scope")

    reset_abbreviation_tracking()
    add_chapter_four_opener(document)
    chapter_four_specs = chapter_four_page_specs()
    for index, spec in enumerate(chapter_four_specs):
        render_chapter_two_page(
            document,
            spec,
            is_last=index == len(chapter_four_specs) - 1,
        )
    add_chapter_four_references(document)

    chapter_five_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_a4(chapter_five_section)
    unlink_section(chapter_five_section)
    set_page_number_format(chapter_five_section, "decimal")
    # No running header — the scope title lives only in the footer.
    add_scope_footer(chapter_five_section, "System Implementation Scope")

    reset_abbreviation_tracking()
    add_chapter_five_opener(document)
    chapter_five_specs = chapter_five_page_specs()
    for index, spec in enumerate(chapter_five_specs):
        render_chapter_two_page(
            document,
            spec,
            is_last=index == len(chapter_five_specs) - 1,
        )
    add_chapter_five_references(document)

    chapter_six_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_a4(chapter_six_section)
    unlink_section(chapter_six_section)
    set_page_number_format(chapter_six_section, "decimal")
    # No running header — the scope title lives only in the footer.
    add_scope_footer(chapter_six_section, "Testing and Evaluation Scope")

    reset_abbreviation_tracking()
    add_chapter_six_opener(document)
    chapter_six_specs = chapter_six_page_specs()
    for index, spec in enumerate(chapter_six_specs):
        render_chapter_two_page(
            document,
            spec,
            is_last=index == len(chapter_six_specs) - 1,
        )
    add_chapter_six_references(document)

    chapter_seven_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_a4(chapter_seven_section)
    unlink_section(chapter_seven_section)
    set_page_number_format(chapter_seven_section, "decimal")
    # No running header — the scope title lives only in the footer.
    add_scope_footer(chapter_seven_section, "Conclusion and Future Work Scope")

    reset_abbreviation_tracking()
    add_chapter_seven_opener(document)
    chapter_seven_specs = chapter_seven_page_specs()
    for index, spec in enumerate(chapter_seven_specs):
        render_chapter_two_page(
            document,
            spec,
            is_last=index == len(chapter_seven_specs) - 1,
        )
    add_chapter_seven_references(document)

    # Whole book prints in black & white (the cover table is the only colour exception).
    monochrome_book_body(document)
    increase_body_content_font_sizes(document)

    document.save(OUTPUT)
    print(f"Created: {OUTPUT}")
    print(f"Chapter One planned pages: 25")
    print(f"Chapter Two planned pages: {1 + len(chapter_two_specs)}")
    print(f"Chapter Three planned pages: {1 + len(chapter_three_specs)}")
    print(f"Chapter Four planned pages: {2 + len(chapter_four_specs)}")
    print(f"Chapter Five planned pages: {2 + len(chapter_five_specs)}")
    print(f"Chapter Six planned pages: {2 + len(chapter_six_specs)}")
    print(f"Chapter Seven planned pages: {2 + len(chapter_seven_specs)}")
    print(f"Inline abbreviation definitions available: {len(ABBREVIATION_DEFINITIONS)}")
    print("Standalone abbreviation pages: disabled")
    print(f"Editorial rules active: {len(EDITORIAL_RULES)}")
    print(f"Figure placeholders: {len(ALL_FIGURES)}")
    print(f"Table entries: {len(ALL_TABLES)}")
    print(
        "IEEE references: "
        f"{len(REFERENCES) + len(CHAPTER_TWO_REFERENCES) + len(CHAPTER_THREE_REFERENCES) + len(CHAPTER_FOUR_REFERENCES) + len(CHAPTER_FIVE_REFERENCES) + len(CHAPTER_SIX_REFERENCES) + len(CHAPTER_SEVEN_REFERENCES)}"
    )


if __name__ == "__main__":
    build()
